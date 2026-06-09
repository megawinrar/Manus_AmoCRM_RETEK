#!/usr/bin/env python3
"""
extract_and_classify.py — Быстрая экстракция + парсинг тендерных файлов.

Цель: за ~3 секунды извлечь текст из всех файлов папки тендера
и распарсить ключевые поля (заказчик, НМЦ, дедлайн, площадка, номер, позиции).

Выход: JSON с полями, готовыми для передачи в emulate_llm_with_dedup.py
или для валидации человеком перед созданием сделки в amoCRM.

Использование:
    python3 scripts/extract_and_classify.py /path/to/tender/folder
    python3 scripts/extract_and_classify.py --files file1.pdf file2.docx file3.xlsx
    python3 scripts/extract_and_classify.py --files file1.pdf file2.docx --json  # только JSON
"""

import argparse
import json
import logging
import os
import re
import subprocess
import sys
import time
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════
# ЭКСТРАКЦИЯ ТЕКСТА
# ═══════════════════════════════════════════════════════════════════

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".rtf", ".txt", ".csv"}


def extract_pdf(filepath: str) -> str:
    """Извлечь текст из PDF через pdftotext (poppler). ~100ms на 1MB."""
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", filepath, "-"],
            capture_output=True, text=True, timeout=30
        )
        return result.stdout
    except Exception as e:
        return f"[ОШИБКА PDF: {e}]"


def is_scanned_pdf(filepath: str, text: str) -> bool:
    """
    Определить, является ли PDF сканированным (без текстового слоя).
    Критерий: < 100 символов на страницу = скорее всего скан/картинка.
    """
    try:
        # Подсчитать страницы через pdfinfo
        info = subprocess.run(
            ["pdfinfo", filepath], capture_output=True, text=True, timeout=10
        )
        pages = 1
        for line in info.stdout.splitlines():
            if line.startswith("Pages:"):
                pages = int(line.split(":")[1].strip())
                break
        chars_per_page = len(text.strip()) / max(pages, 1)
        return chars_per_page < 100
    except Exception:
        # Fallback: если текст почти пустой
        return len(text.strip()) < 200


def ocr_pdf(filepath: str) -> str:
    """
    OCR для сканированных PDF: pdf2image → pytesseract.
    Используется ТОЛЬКО когда pdftotext не дал результата.
    Время: ~3-5 сек на страницу.
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract

        logger.info(f"OCR fallback для: {os.path.basename(filepath)}")
        images = convert_from_path(filepath, dpi=300)
        text_parts = []
        for i, img in enumerate(images):
            page_text = pytesseract.image_to_string(img, lang='rus+eng')
            text_parts.append(page_text)
            logger.debug(f"  OCR стр.{i+1}: {len(page_text)} символов")
        return "\n\n".join(text_parts)
    except ImportError:
        return "[OCR НЕДОСТУПЕН: установите pytesseract и pdf2image]"
    except Exception as e:
        return f"[ОШИБКА OCR: {e}]"


def ocr_region_around_keyword(filepath: str, keyword: str, margin_px: int = 200) -> str:
    """
    Точечный OCR: найти область вокруг ключевого слова и распознать её.
    Используется для верификации конкретного поля когда confidence низкая.
    
    Стратегия:
    1. Получить координаты слова через pdftotext -bbox
    2. Вырезать регион (keyword bbox + margin)
    3. OCR только этого региона
    
    Возвращает распознанный текст региона или пустую строку.
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract
        from PIL import Image

        # Получить bbox всех слов
        result = subprocess.run(
            ["pdftotext", "-bbox", filepath, "-"],
            capture_output=True, text=True, timeout=30
        )

        # Найти координаты ключевого слова
        keyword_lower = keyword.lower()
        bbox_match = None
        page_num = 0

        for line in result.stdout.splitlines():
            if '<page' in line:
                page_num += 1
            if keyword_lower in line.lower() and 'xMin' in line:
                # Парсим bbox из HTML-подобного формата pdftotext -bbox
                coords = re.findall(r'[xy](?:Min|Max)="([\d.]+)"', line)
                if len(coords) >= 4:
                    bbox_match = {
                        'page': page_num,
                        'xMin': float(coords[0]),
                        'yMin': float(coords[1]),
                        'xMax': float(coords[2]),
                        'yMax': float(coords[3]),
                    }
                    break

        if not bbox_match:
            return ""

        # Конвертировать нужную страницу в изображение
        images = convert_from_path(
            filepath, dpi=300,
            first_page=bbox_match['page'],
            last_page=bbox_match['page']
        )
        if not images:
            return ""

        img = images[0]
        w, h = img.size

        # Масштабировать координаты (PDF points → pixels at 300dpi)
        scale = 300 / 72  # PDF uses 72 dpi
        x1 = max(0, int(bbox_match['xMin'] * scale) - margin_px)
        y1 = max(0, int(bbox_match['yMin'] * scale) - margin_px)
        x2 = min(w, int(bbox_match['xMax'] * scale) + margin_px)
        y2 = min(h, int(bbox_match['yMax'] * scale) + margin_px)

        # Вырезать и OCR
        region = img.crop((x1, y1, x2, y2))
        text = pytesseract.image_to_string(region, lang='rus+eng')
        return text.strip()

    except ImportError:
        return ""
    except Exception as e:
        logger.warning(f"OCR region failed for '{keyword}': {e}")
        return ""


def extract_docx(filepath: str) -> str:
    """Извлечь текст из DOCX через python-docx. ~300ms на 2.5MB."""
    try:
        from docx import Document
        doc = Document(filepath)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("\t".join(cells))
        return "\n".join(parts)
    except Exception as e:
        return f"[ОШИБКА DOCX: {e}]"


def extract_xlsx(filepath: str) -> str:
    """Извлечь данные из XLSX через openpyxl. ~700ms."""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
        parts = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            parts.append(f"=== Лист: {sheet_name} ===")
            for row in ws.iter_rows(values_only=True):
                vals = [str(v) if v is not None else "" for v in row]
                line = "\t".join(vals)
                if line.strip():
                    parts.append(line)
        wb.close()
        return "\n".join(parts)
    except Exception as e:
        return f"[ОШИБКА XLSX: {e}]"


def extract_doc(filepath: str) -> str:
    """Извлечь текст из .doc через antiword или catdoc."""
    for tool in ["antiword", "catdoc"]:
        try:
            result = subprocess.run(
                [tool, filepath], capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout
        except FileNotFoundError:
            continue
    return "[ОШИБКА DOC: antiword/catdoc не установлены]"


def extract_file(filepath: str) -> tuple:
    """
    Извлечь текст из файла. Возвращает (text, time_ms, method).
    method: 'pdftotext' | 'ocr' | 'python-docx' | 'openpyxl' | 'text' | 'antiword'
    """
    ext = Path(filepath).suffix.lower()
    t0 = time.time()
    method = "unknown"

    if ext == ".pdf":
        text = extract_pdf(filepath)
        method = "pdftotext"
        # Проверка: если PDF сканированный — fallback на OCR
        if is_scanned_pdf(filepath, text):
            logger.info(f"PDF сканированный, применяю OCR: {os.path.basename(filepath)}")
            text = ocr_pdf(filepath)
            method = "ocr"
    elif ext == ".docx":
        text = extract_docx(filepath)
        method = "python-docx"
    elif ext == ".xlsx" or ext == ".xls":
        text = extract_xlsx(filepath)
        method = "openpyxl"
    elif ext == ".doc":
        text = extract_doc(filepath)
        method = "antiword"
    elif ext in (".txt", ".csv", ".rtf"):
        try:
            text = Path(filepath).read_text(encoding="utf-8", errors="replace")
            method = "text"
        except Exception as e:
            text = f"[ОШИБКА: {e}]"
    else:
        text = ""

    elapsed_ms = (time.time() - t0) * 1000
    return text, elapsed_ms, method


# ═══════════════════════════════════════════════════════════════════
# ПАРСИНГ КЛЮЧЕВЫХ ПОЛЕЙ (regex + эвристики)
# ═══════════════════════════════════════════════════════════════════

def parse_tender_fields(full_text: str, file_names: list) -> dict:
    """
    Извлечь ключевые поля тендера из объединённого текста.
    Возвращает dict с полями, совместимыми с emulate_llm_with_dedup.py.
    """
    result = {
        "customer": None,
        "customer_short": None,
        "nmc": None,
        "nmc_formatted": None,
        "deadline": None,
        "deadline_raw": None,
        "platform": None,
        "platform_url": None,
        "procedure_number": None,
        "notice_number": None,
        "procedure_type": None,
        "subject": None,
        "position_count": None,
        "total_quantity": None,
        "equivalent_allowed": None,
        "gisp_required": None,
        "delivery_terms": None,
        "city": None,
        "direction_hint": None,
        "situation_hint": None,
        "priority_hint": None,
        "key_products": [],
        "warnings": [],
    }

    # --- ЗАКАЗЧИК ---
    # Стратегия: ищем организационно-правовую форму + название в кавычках
    # Приоритет: 1) Акционерное общество/ООО/ФГУП + полное имя
    #            2) АО/ПАО/ООО + «...»
    #            3) Поле "Заказчик:" / "Организатор:"
    customer_patterns = [
        # Акционерное общество "..." или „..."
        r'((?:Акционерное общество|Публичное акционерное общество|Общество с ограниченной ответственностью|'
        r'Федеральное государственное унитарное предприятие|Государственная корпорация)'
        r'\s*[„«"\'\u201e\u201c].*?["\u201d»\'\u201f])',
        # АО/ПАО/ООО/ФГУП + «...» или "..."
        r'((?:АО|ПАО|ООО|ФГУП|ФКП|ГК)\s*[«„"\'\u201e\u201c][^»"\n]{3,80}[»"\u201d\u201f\'\u201f])',
        # Наименование + значение (табличный формат площадок)
        r'Наименование\s+(?!заказа)(?!закупки)\s*((?:Акционерное|Публичное|Общество|Федеральное|Государственн)[^\n]{10,150})',
        # "Заказчик: ..." или "Организатор: ..."
        r'(?:Заказчик|Организатор|Покупатель)\s*[:\-—]\s*(.+?)(?:\n|$)',
    ]
    for pat in customer_patterns:
        m = re.search(pat, full_text)
        if m:
            customer = m.group(1).strip()
            # Фильтр: не должен содержать "поставка", "инструмент" (это предмет, не заказчик)
            if (len(customer) > 5 and len(customer) < 200
                    and "поставка" not in customer.lower()
                    and "инструмент" not in customer.lower()):
                result["customer"] = customer
                # Сокращённое имя
                short = re.sub(
                    r'^(Акционерное общество|Публичное акционерное общество|'
                    r'Общество с ограниченной ответственностью|'
                    r'Федеральное государственное унитарное предприятие|'
                    r'Государственная корпорация|АО|ПАО|ООО|ФГУП|ФКП|ГК)\s*',
                    '', customer, flags=re.IGNORECASE
                )
                short = re.sub(r'[«»„""\u201e\u201c\u201d\u201f\']+', '', short).strip()
                result["customer_short"] = short[:50]
                break

    # --- НМЦ (Начальная максимальная цена) ---
    nmc_patterns = [
        r'(?:НМЦ|Начальная.*?максимальная.*?цена|Цена договора|Общая стоимость)\s*[:\-—]?\s*([\d\s.,]+)\s*(?:руб|₽|RUB)',
        r'(?:Итого|ИТОГО|Всего)\s*(?:с\s*НДС)?\s*[:\-—]?\s*([\d\s.,]+)\s*(?:руб|₽)',
        r'(?:Максимальная цена)\s*[:\-—]?\s*([\d\s.,]+)',
    ]
    for pat in nmc_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            nmc_str = m.group(1).strip()
            nmc_str = re.sub(r'\s+', '', nmc_str)
            nmc_str = nmc_str.replace(',', '.')
            # Убрать лишние точки (тысячные разделители)
            parts = nmc_str.split('.')
            if len(parts) > 2:
                # Последняя часть — копейки
                nmc_str = ''.join(parts[:-1]) + '.' + parts[-1]
            try:
                nmc = float(nmc_str)
                if nmc > 1000:  # Минимальный порог
                    result["nmc"] = nmc
                    result["nmc_formatted"] = f"{nmc:,.2f} руб.".replace(',', ' ')
                    break
            except ValueError:
                continue

    # --- ДЕДЛАЙН (Срок подачи заявок) ---
    deadline_patterns = [
        r'(?:Окончание.*?подачи|Срок подачи|Дата окончания подачи|Подача.*?до)\s*[:\-—]?\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})\s*(\d{1,2}:\d{2})?',
        r'(?:Окончание срока подачи заявок)\s+(\d{1,2}\.\d{1,2}\.\d{4})\s+(\d{1,2}:\d{2})',
    ]
    for pat in deadline_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            date_str = m.group(1)
            time_str = m.group(2) if m.lastindex >= 2 and m.group(2) else ""
            result["deadline_raw"] = f"{date_str} {time_str}".strip()
            # Нормализация даты
            date_str = date_str.replace('/', '.')
            parts = date_str.split('.')
            if len(parts) == 3:
                d, mo, y = parts
                if len(y) == 2:
                    y = "20" + y
                result["deadline"] = f"{y}-{mo.zfill(2)}-{d.zfill(2)}"
            break

    # --- ПЛОЩАДКА ---
    platform_patterns = [
        r'(?:www\.)?(etprf\.ru|b2b-center\.ru|zakupki\.gov\.ru|roseltorg\.ru|'
        r'etp-ets\.ru|fabrikant\.ru|rts-tender\.ru|sberbank-ast\.ru|'
        r'lot-online\.ru|tektorg\.ru|astgoz\.ru|etp-micex\.ru)',
        r'(?:Электронная торговая площадка|ЭТП|Площадка)\s*[:\-—]?\s*[«"]?(.+?)[»"]?\s*(?:\n|$)',
    ]
    for pat in platform_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            platform = m.group(1).strip() if m.lastindex >= 1 else m.group(0)
            # Нормализация
            platform_map = {
                "etprf.ru": ("ЕТПРФ", "https://www.etprf.ru"),
                "b2b-center.ru": ("B2B-Center", "https://www.b2b-center.ru"),
                "zakupki.gov.ru": ("Госзакупки (ЕИС)", "https://zakupki.gov.ru"),
                "roseltorg.ru": ("Росэлторг", "https://www.roseltorg.ru"),
                "etp-ets.ru": ("ЕТС", "https://etp-ets.ru"),
                "fabrikant.ru": ("Фабрикант", "https://www.fabrikant.ru"),
                "rts-tender.ru": ("РТС-тендер", "https://www.rts-tender.ru"),
                "sberbank-ast.ru": ("Сбербанк-АСТ", "https://www.sberbank-ast.ru"),
                "lot-online.ru": ("Lot-Online", "https://www.lot-online.ru"),
                "tektorg.ru": ("ТЭК-Торг", "https://www.tektorg.ru"),
                "astgoz.ru": ("АСТ ГОЗ", "https://www.astgoz.ru"),
            }
            for domain, (name, url) in platform_map.items():
                if domain in platform.lower():
                    result["platform"] = name
                    result["platform_url"] = url
                    break
            if not result["platform"]:
                result["platform"] = platform[:50]
            break

    # --- НОМЕР ПРОЦЕДУРЫ ---
    proc_patterns = [
        r'(?:Номер закупки|Реестровый номер|Номер процедуры|Номер извещения)\s*[:\-—]?\s*(\d[\d\-./]+\d)',
        r'(\d{4}-\d{4}-\d{5})',  # формат ХХХХ-ХХХХ-ХХХХХ
        r'(?:извещени[еяю])\s*(?:№|#)?\s*(\d{10,})',  # номер извещения (10+ цифр)
    ]
    for pat in proc_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            num = m.group(1).strip()
            if re.match(r'\d{4}-\d{4}-\d{5}', num):
                result["procedure_number"] = num
            elif len(num) >= 10 and num.isdigit():
                result["notice_number"] = num
            else:
                result["procedure_number"] = num
            break

    # Отдельно ищем номер извещения если не нашли
    if not result["notice_number"]:
        m = re.search(r'(?:извещени[еяю]|Извещение)\s*(?:№|#|номер)?\s*(\d{10,})', full_text)
        if m:
            result["notice_number"] = m.group(1)

    # Отдельно ищем номер закупки формата ХХХХ-ХХХХ-ХХХХХ
    if not result["procedure_number"]:
        m = re.search(r'(\d{4}-\d{4}-\d{5})', full_text)
        if m:
            result["procedure_number"] = m.group(1)

    # --- ТИП ПРОЦЕДУРЫ ---
    proc_type_patterns = [
        r'(?:Способ закупки|Вид процедуры|Тип закупки|Способ определения)\s*[:\-—]?\s*(.+?)(?:\n|$)',
        r'(Закрытый\s+(?:запрос\s+котировок|аукцион|конкурс|редукцион))',
        r'(Открытый\s+(?:запрос\s+котировок|аукцион|конкурс))',
        r'(Запрос\s+котировок)',
        r'(Запрос\s+предложений)',
    ]
    for pat in proc_type_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            result["procedure_type"] = m.group(1).strip()[:80]
            break

    # --- ПРЕДМЕТ ЗАКУПКИ ---
    subject_patterns = [
        # "Наименование заказа" (формат площадок)
        r'Наименование заказа\s+(.+?)(?:\n|$)',
        # "Предмет договора" (без префикса "Полное наименование")
        r'(?<!Полное наименование \()предмет договора[)\s]*[:\-—]?\s*(.+?)(?:\n|$)',
        # "Предмет закупки" / "Наименование закупки" / "Объект закупки"
        r'(?:Предмет|Наименование закупки|Объект закупки)\s*[:\-—]?\s*(.+?)(?:\n|$)',
    ]
    for pat in subject_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            subj = m.group(1).strip()
            # Фильтр: не должен быть слишком коротким и не должен быть номером
            if len(subj) > 10 and not subj[0].isdigit():
                result["subject"] = subj[:200]
                break
            elif len(subj) > 10:
                # Если начинается с номера закупки, убрать его
                subj_clean = re.sub(r'^\d{4}-\d{4}-\d{5}[.\s]*', '', subj).strip()
                if len(subj_clean) > 10:
                    result["subject"] = subj_clean[:200]
                    break
                else:
                    result["subject"] = subj[:200]
                    break

    # Fallback: из имени файла
    if not result["subject"]:
        for fn in file_names:
            if "Поставка" in fn or "поставка" in fn:
                subj = re.sub(r'[_—\-]+', ' ', fn)
                subj = re.sub(r'\.\w+$', '', subj)
                subj = re.sub(r'^\d{4}-\d{4}-\d{5}[.\s]*', '', subj).strip()
                result["subject"] = subj[:200]
                break

    # --- КОЛИЧЕСТВО ПОЗИЦИЙ ---
    # Ищем таблицы с нумерацией
    position_numbers = re.findall(r'^\s*(\d{1,3})\s*[.\t|]', full_text, re.MULTILINE)
    if position_numbers:
        max_pos = max(int(n) for n in position_numbers if int(n) <= 500)
        if max_pos >= 2:
            result["position_count"] = max_pos

    # --- ЭКВИВАЛЕНТ ---
    # Ищем явное "Допускается эквивалент: Да/Нет" (формат площадок)
    equiv_explicit = re.search(r'Допускается\s+эквивалент\s+(Да|Нет)', full_text, re.IGNORECASE)
    if equiv_explicit:
        result["equivalent_allowed"] = equiv_explicit.group(1).lower() == "да"
    elif re.search(r'(?:или\s+эквивалент|эквивалент\s*допускается|аналог\s*допускается)', full_text, re.IGNORECASE):
        result["equivalent_allowed"] = True
    elif re.search(r'(?:без\s*эквивалент|эквивалент\s*не\s*допускается)', full_text, re.IGNORECASE):
        result["equivalent_allowed"] = False

    # --- ГИСП ---
    if re.search(r'(?:ГИСП|gisp\.gov\.ru|ЕЭТП|реестр\s*промышленной\s*продукции)', full_text, re.IGNORECASE):
        result["gisp_required"] = True

    # --- ГОРОД ---
    city_patterns = [
        r'(?:г\.\s*|город\s+)(Москв[аы]|Санкт-Петербург[а]?|Нижн[а-я]*\s*Новгород[а]?|'
        r'Екатеринбург[а]?|Новосибирск[а]?|Казан[ьи]|Челябинск[а]?|Омск[а]?|'
        r'Самар[аы]|Ростов[а]?(?:\s*-на-Дону)?|Уф[аы]|Красноярск[а]?|'
        r'Пермь|Перми|Воронеж[а]?|Волгоград[а]?|Тул[аы]|Ижевск[а]?|'
        r'Курган[а]?|Ковров[а]?|[А-ЯЁ][а-яё]+)',
    ]
    for pat in city_patterns:
        m = re.search(pat, full_text)
        if m:
            city = m.group(1).strip()
            if len(city) > 2:
                result["city"] = f"г. {city}"
                break

    # --- НАПРАВЛЕНИЕ (эвристика по ключевым словам) ---
    text_lower = full_text.lower()

    # Подсчёт ключевых слов для определения направления
    direction_scores = {
        "SPEC-DRAWING": 0,
        "HSS-STANDARD": 0,
        "CARBIDE-STANDARD": 0,
        "DIAMOND-STANDARD": 0,
        "OUT-OF-SCOPE": 0,
    }

    # SPEC-DRAWING keywords
    for kw in ["по чертеж", "по тз", "по эскиз", "по образц", "нестандартн", "специнструмент"]:
        direction_scores["SPEC-DRAWING"] += len(re.findall(kw, text_lower))

    # HSS-STANDARD keywords
    for kw in ["гост", "метчик", "сверло ", "развертка", "развёртка", "плашка", "зенкер",
               "hss", "р6м5", "р18", "быстрорез"]:
        direction_scores["HSS-STANDARD"] += len(re.findall(kw, text_lower))

    # CARBIDE-STANDARD keywords
    for kw in ["твердосплав", "фреза ", "фрезы ", "пластин", "carbide", "vhm",
               "концев", "торцев", "гравер", "promatool", "gesac", "korloy",
               "sandvik", "iscar", "kennametal", "mitsubishi", "tungaloy"]:
        direction_scores["CARBIDE-STANDARD"] += len(re.findall(kw, text_lower))

    # DIAMOND keywords
    for kw in ["алмаз", "diamond", "pcd", "cbn", "кубический нитрид бора"]:
        direction_scores["DIAMOND-STANDARD"] += len(re.findall(kw, text_lower))

    # OUT-OF-SCOPE keywords
    for kw in ["калибр", "скоба", "штангенциркуль", "микрометр", "индикатор"]:
        direction_scores["OUT-OF-SCOPE"] += len(re.findall(kw, text_lower))

    # Определить направление по максимальному score
    max_dir = max(direction_scores, key=direction_scores.get)
    max_score = direction_scores[max_dir]
    if max_score > 0:
        result["direction_hint"] = max_dir
        # Если два направления конкурируют — предупредить
        sorted_dirs = sorted(direction_scores.items(), key=lambda x: -x[1])
        if len(sorted_dirs) > 1 and sorted_dirs[1][1] > 0:
            ratio = sorted_dirs[1][1] / max(sorted_dirs[0][1], 1)
            if ratio > 0.5:
                result["warnings"].append(
                    f"Конкурирующие направления: {sorted_dirs[0][0]}({sorted_dirs[0][1]}) "
                    f"vs {sorted_dirs[1][0]}({sorted_dirs[1][1]})"
                )

    # --- ТИП СИТУАЦИИ (эвристика) ---
    has_ioz = any("иоз" in fn.lower() or "идоз" in fn.lower() for fn in file_names)
    has_tz = any("техническое задание" in fn.lower() or "тз" in fn.lower() for fn in file_names)
    has_spec = any("спецификаци" in fn.lower() for fn in file_names)
    has_contract = any("договор" in fn.lower() or "контракт" in fn.lower() for fn in file_names)
    has_nmc_file = any("нмц" in fn.lower() or "обоснование" in fn.lower() for fn in file_names)

    if has_ioz or has_tz or has_contract or has_nmc_file:
        result["situation_hint"] = "Запрос котировок / реальные торги"
    elif "соз" in " ".join(file_names).lower() or "сбор" in text_lower[:500]:
        result["situation_hint"] = "СОЗ"
    else:
        result["situation_hint"] = "Неясно"

    # --- ПРИОРИТЕТ (эвристика) ---
    nmc = result.get("nmc", 0) or 0
    if nmc > 20_000_000:
        result["priority_hint"] = "Р2"  # Крупная сумма
    elif nmc > 5_000_000:
        result["priority_hint"] = "Р2"  # Средняя+
    elif nmc > 1_000_000:
        result["priority_hint"] = "Р3"
    elif nmc > 0:
        result["priority_hint"] = "Р3"
    else:
        result["priority_hint"] = "Р3"  # Default

    # Если дедлайн ≤ 48ч — автоэскалация до Р1 (будет сделана post-create hook)
    if result.get("deadline"):
        from datetime import datetime, timedelta
        try:
            dl = datetime.strptime(result["deadline"], "%Y-%m-%d")
            now = datetime.now()
            hours_left = (dl - now).total_seconds() / 3600
            if hours_left <= 48:
                result["priority_hint"] = "Р1"
                result["warnings"].append(f"СРОЧНО: до дедлайна {hours_left:.0f}ч (≤48ч → Р1)")
            elif hours_left <= 120:
                result["warnings"].append(f"Внимание: до дедлайна {hours_left:.0f}ч (~{hours_left/24:.0f} дней)")
        except ValueError:
            pass

    # --- КЛЮЧЕВЫЕ ПРОДУКТЫ (первые 10 уникальных) ---
    product_keywords = set()
    for kw in re.findall(r'(?:фреза|фрезы|метчик|сверло|развертка|развёртка|пластина|гравер|'
                         r'зенкер|плашка|резец|борфреза|державка|патрон|оправка)\w*',
                         text_lower):
        product_keywords.add(kw[:20])
    result["key_products"] = sorted(list(product_keywords))[:10]

    # --- СРОКИ ПОСТАВКИ ---
    delivery_patterns = [
        r'(?:Срок поставки|Срок исполнения)\s*[:\-—]?\s*(.+?)(?:\n|$)',
        r'(\d+)\s*(?:календарных|рабочих)?\s*дн[а-я]*\s*(?:с момента|после|от даты)',
        r'(?:партиями|(?:не позднее|до)\s+\d{1,2}[./]\d{1,2}[./]\d{2,4})',
    ]
    for pat in delivery_patterns:
        m = re.search(pat, full_text, re.IGNORECASE)
        if m:
            delivery = m.group(0).strip()
            # Фильтр: не должен быть именем файла
            if not delivery.endswith(('.pdf', '.docx', '.xlsx', '.doc')):
                result["delivery_terms"] = delivery[:100]
                break

    return result


# ═══════════════════════════════════════════════════════════════════
# ФОРМАТИРОВАНИЕ РЕЗУЛЬТАТА
# ═══════════════════════════════════════════════════════════════════

def format_result(parsed: dict, extraction_stats: list, total_time: float) -> str:
    """Форматировать результат для вывода в терминал."""
    lines = []
    lines.append("=" * 70)
    lines.append("  РЕЗУЛЬТАТ ЭКСТРАКЦИИ И ПАРСИНГА ТЕНДЕРА")
    lines.append("=" * 70)
    lines.append("")

    # Статистика экстракции
    lines.append("  📁 Экстракция файлов:")
    for stat in extraction_stats:
        lines.append(f"    {stat['ext']:6s} | {stat['time_ms']:6.0f} ms | "
                     f"{stat['size_kb']:7.0f} KB | {stat['chars']:>7,} chars | {stat['name'][:45]}")
    lines.append(f"  ─── Итого: {total_time*1000:.0f} ms ({total_time:.2f} сек) ───")
    lines.append("")

    # Распознанные поля
    lines.append("  📋 Распознанные поля:")
    lines.append(f"    Заказчик:        {parsed.get('customer') or '❌ НЕ НАЙДЕН'}")
    lines.append(f"    Заказчик (кратко): {parsed.get('customer_short') or '—'}")
    lines.append(f"    НМЦ:             {parsed.get('nmc_formatted') or '❌ НЕ НАЙДЕНА'}")
    lines.append(f"    Дедлайн подачи:  {parsed.get('deadline') or '❌ НЕ НАЙДЕН'} ({parsed.get('deadline_raw', '')})")
    lines.append(f"    Площадка:        {parsed.get('platform') or '—'} ({parsed.get('platform_url', '')})")
    lines.append(f"    Номер закупки:   {parsed.get('procedure_number') or '—'}")
    lines.append(f"    Номер извещения: {parsed.get('notice_number') or '—'}")
    lines.append(f"    Тип процедуры:   {parsed.get('procedure_type') or '—'}")
    lines.append(f"    Предмет:         {parsed.get('subject') or '—'}")
    lines.append(f"    Кол-во позиций:  {parsed.get('position_count') or '—'}")
    lines.append(f"    Город:           {parsed.get('city') or '—'}")
    lines.append(f"    Эквивалент:      {parsed.get('equivalent_allowed')}")
    lines.append(f"    ГИСП:            {parsed.get('gisp_required')}")
    lines.append(f"    Поставка:        {parsed.get('delivery_terms') or '—'}")
    lines.append("")

    # Классификация (подсказки)
    lines.append("  🎯 Подсказки классификации:")
    lines.append(f"    Направление:     {parsed.get('direction_hint') or '—'}")
    lines.append(f"    Тип ситуации:    {parsed.get('situation_hint') or '—'}")
    lines.append(f"    Приоритет:       {parsed.get('priority_hint') or '—'}")
    lines.append(f"    Продукты:        {', '.join(parsed.get('key_products', []))}")
    lines.append("")

    # Предупреждения
    if parsed.get("warnings"):
        lines.append("  ⚠️  Предупреждения:")
        for w in parsed["warnings"]:
            lines.append(f"    • {w}")
        lines.append("")

    lines.append("=" * 70)
    return "\n".join(lines)


def build_emulate_command(parsed: dict, files: list, tender_path: str = "") -> str:
    """Сгенерировать команду для emulate_llm_with_dedup.py."""
    parts = ['python3 scripts/emulate_llm_with_dedup.py']
    parts.append('  --files \\')
    for f in files:
        parts.append(f'    "{f}" \\')

    if tender_path:
        parts.append(f'  --tender-path "{tender_path}" \\')
    if parsed.get("customer"):
        parts.append(f'  --customer "{parsed["customer"]}" \\')
    if parsed.get("nmc"):
        parts.append(f'  --nmc {parsed["nmc"]} \\')
    if parsed.get("priority_hint"):
        parts.append(f'  --priority {parsed["priority_hint"]} \\')
    if parsed.get("direction_hint"):
        parts.append(f'  --direction {parsed["direction_hint"]} \\')
    if parsed.get("situation_hint"):
        parts.append(f'  --situation "{parsed["situation_hint"]}" \\')
    if parsed.get("procedure_type"):
        parts.append(f'  --procedure-type "{parsed["procedure_type"]}" \\')
    if parsed.get("procedure_number"):
        parts.append(f'  --procedure-number "{parsed["procedure_number"]}" \\')
    if parsed.get("position_count"):
        parts.append(f'  --positions {parsed["position_count"]} \\')
    if parsed.get("deadline"):
        parts.append(f'  --deadline "{parsed["deadline"]}" \\')

    # Comment
    comment_parts = []
    if parsed.get("position_count"):
        comment_parts.append(f"{parsed['position_count']} позиций")
    if parsed.get("key_products"):
        comment_parts.append(f"({', '.join(parsed['key_products'][:5])})")
    if parsed.get("platform"):
        comment_parts.append(f"Площадка: {parsed['platform']}")
    if parsed.get("gisp_required"):
        comment_parts.append("Требуется ГИСП")
    if parsed.get("equivalent_allowed"):
        comment_parts.append("Эквивалент допускается")

    if comment_parts:
        parts.append(f'  --comment "{". ".join(comment_parts)}" \\')

    parts.append('  --confidence 0.90 \\')
    parts.append('  --source "Тендерная площадка"')

    return "\n".join(parts)




# ═══════════════════════════════════════════════════════════════════
# ВАЛИДАЦИЯ + OCR-FALLBACK
# ═══════════════════════════════════════════════════════════════════

# Все поля, которые идут в карточку amoCRM, с ключевыми словами для OCR-верификации
# severity: "block" = без этого нельзя создать сделку, "warn" = создать можно, но пометить
REQUIRED_FIELDS = {
    # --- Критичные (block) ---
    "customer": {
        "label": "Заказчик",
        "severity": "block",
        "ocr_keywords": ["Наименование", "Заказчик", "Организатор"],
    },
    "nmc": {
        "label": "НМЦ",
        "severity": "block",
        "ocr_keywords": ["НМЦ", "максимальная", "Итого"],
    },
    "deadline": {
        "label": "Дедлайн",
        "severity": "block",
        "ocr_keywords": ["подачи", "Окончание", "срок"],
    },
    # --- Важные (warn) ---
    "procedure_number": {
        "label": "Номер закупки",
        "severity": "warn",
        "ocr_keywords": ["Номер", "закупки", "извещения"],
    },
    "platform": {
        "label": "Площадка",
        "severity": "warn",
        "ocr_keywords": ["площадка", "ЭТП", "etprf"],
    },
    "position_count": {
        "label": "Кол-во позиций",
        "severity": "warn",
        "ocr_keywords": ["позиц", "Лот", "Спецификац"],
    },
    "equivalent_allowed": {
        "label": "Эквивалент",
        "severity": "warn",
        "ocr_keywords": ["эквивалент", "аналог", "или эквивалент"],
    },
    "gisp_required": {
        "label": "ГИСП",
        "severity": "warn",
        "ocr_keywords": ["ГИСП", "gisp", "реестр"],
    },
    "procedure_type": {
        "label": "Тип процедуры",
        "severity": "warn",
        "ocr_keywords": ["способ", "закупки", "котировок", "аукцион"],
    },
    "subject": {
        "label": "Предмет закупки",
        "severity": "warn",
        "ocr_keywords": ["Предмет", "Наименование закупки", "Поставка"],
    },
    "city": {
        "label": "Город",
        "severity": "warn",
        "ocr_keywords": ["Место", "адрес", "поставки"],
    },
    "notice_number": {
        "label": "Номер извещения",
        "severity": "warn",
        "ocr_keywords": ["извещени", "Номер", "реестровый"],
    },
    "direction_hint": {
        "label": "Направление",
        "severity": "warn",
        "ocr_keywords": [],  # определяется по продуктам, OCR не нужен
    },
}

# Поля, где низкая уверенность требует OCR-проверки
CONFIDENCE_THRESHOLDS = {
    "nmc": {"min_value": 100_000, "max_value": 10_000_000_000},
    "deadline": {"min_year": 2024, "max_year": 2030},
    "position_count": {"min_value": 1, "max_value": 5000},
}


def validate_and_ocr_fallback(parsed: dict, pdf_files: list, full_text: str) -> dict:
    """
    Валидация обязательных полей + OCR-fallback.
    
    Логика:
    1. Поле не найдено (None) → OCR вокруг ключевого слова
    2. Поле найдено, но значение сомнительно (confidence < 0.85) → OCR-верификация
    3. OCR не помог → warning + пометка "требует ручной проверки"
    
    Добавляет в parsed:
    - confidence_scores: {field: 0.0-1.0}
    - ocr_applied: [field_names]
    - validation_status: 'ok' | 'warnings' | 'blocked'
    """
    confidence_scores = {}
    ocr_applied = []
    missing_critical = []
    
    for field, config in REQUIRED_FIELDS.items():
        value = parsed.get(field)
        
        # === Поле не найдено ===
        if value is None:
            confidence_scores[field] = 0.0
            
            # Попытка OCR-fallback для PDF
            if pdf_files:
                ocr_text = ""
                for kw in config["ocr_keywords"]:
                    for pdf_path in pdf_files:
                        ocr_text = ocr_region_around_keyword(pdf_path, kw, margin_px=250)
                        if ocr_text and len(ocr_text) > 10:
                            break
                    if ocr_text:
                        break
                
                if ocr_text:
                    # Повторный парсинг на OCR-тексте
                    ocr_parsed = parse_tender_fields(ocr_text, [])
                    if ocr_parsed.get(field) is not None:
                        parsed[field] = ocr_parsed[field]
                        if field == "customer" and ocr_parsed.get("customer_short"):
                            parsed["customer_short"] = ocr_parsed["customer_short"]
                        if field == "nmc" and ocr_parsed.get("nmc_formatted"):
                            parsed["nmc_formatted"] = ocr_parsed["nmc_formatted"]
                        confidence_scores[field] = 0.7
                        ocr_applied.append(field)
                        parsed["warnings"].append(
                            f"Поле '{config['label']}' извлечено через OCR (требует проверки)"
                        )
                        continue
            
            # OCR не помог
            if config["severity"] == "block":
                missing_critical.append(config["label"])
                parsed["warnings"].append(
                    f"КРИТИЧНО: '{config['label']}' не найдено — требуется ручной ввод"
                )
            else:
                parsed["warnings"].append(
                    f"Поле '{config['label']}' не найдено (некритично)"
                )
            continue
        
        # === Поле найдено — проверка адекватности ===
        confidence = 1.0
        
        # Проверка НМЦ: диапазон 100К - 10 млрд
        if field == "nmc" and field in CONFIDENCE_THRESHOLDS:
            thresholds = CONFIDENCE_THRESHOLDS[field]
            nmc_val = parsed["nmc"]
            if nmc_val < thresholds["min_value"]:
                confidence = 0.5
                parsed["warnings"].append(
                    f"НМЦ подозрительно мала: {nmc_val:,.0f} руб. (< {thresholds['min_value']:,.0f})"
                )
            elif nmc_val > thresholds["max_value"]:
                confidence = 0.4
                parsed["warnings"].append(
                    f"НМЦ подозрительно велика: {nmc_val:,.0f} руб. (> {thresholds['max_value']:,.0f})"
                )
        
        # Проверка дедлайна: год 2024-2030
        if field == "deadline" and field in CONFIDENCE_THRESHOLDS:
            thresholds = CONFIDENCE_THRESHOLDS[field]
            try:
                year = int(parsed["deadline"][:4])
                if year < thresholds["min_year"] or year > thresholds["max_year"]:
                    confidence = 0.3
                    parsed["warnings"].append(
                        f"Дедлайн подозрительный: {parsed['deadline']} (год вне диапазона)"
                    )
            except (ValueError, IndexError):
                confidence = 0.5
        
        # Проверка заказчика: орг.форма + длина
        if field == "customer":
            cust = parsed["customer"]
            if len(cust) < 10:
                confidence = 0.5
            elif not any(kw in cust for kw in ["общество", "АО", "ПАО", "ООО", "ФГУП", "ГК", "ФКП",
                                                "Общество", "Акционерное", "Федеральное"]):
                confidence = 0.6
        
        # Проверка кол-ва позиций: диапазон 1-5000
        if field == "position_count" and field in CONFIDENCE_THRESHOLDS:
            thresholds = CONFIDENCE_THRESHOLDS[field]
            pos_val = parsed["position_count"]
            if pos_val < thresholds["min_value"]:
                confidence = 0.4
                parsed["warnings"].append(
                    f"Кол-во позиций = {pos_val} (подозрительно)"
                )
            elif pos_val > thresholds["max_value"]:
                confidence = 0.5
                parsed["warnings"].append(
                    f"Кол-во позиций = {pos_val} (подозрительно много)"
                )
        
        # Проверка эквивалента: должен быть bool
        if field == "equivalent_allowed":
            if not isinstance(value, bool):
                confidence = 0.6
        
        # Проверка ГИСП: должен быть bool
        if field == "gisp_required":
            if not isinstance(value, bool):
                confidence = 0.6
        
        # Проверка типа процедуры: должен содержать известные ключевые слова
        if field == "procedure_type":
            known_types = ["запрос котировок", "аукцион", "конкурс", "запрос предложений",
                           "единственный поставщик", "открытый", "закрытый"]
            if not any(kt in str(value).lower() for kt in known_types):
                confidence = 0.6
                parsed["warnings"].append(
                    f"Тип процедуры не распознан как стандартный: '{value}'"
                )
        
        # Проверка предмета: должен содержать слова про инструмент/поставку
        if field == "subject":
            if len(str(value)) < 10:
                confidence = 0.5
            elif not any(kw in str(value).lower() for kw in ["поставка", "инструмент", "выполнение",
                                                             "оказание", "изготовление", "производство"]):
                confidence = 0.7  # не снижаем сильно, может быть нестандартная формулировка
        
        # Проверка города: должен содержать "г." или известный город
        if field == "city":
            known_cities = ["Москв", "Петербург", "Новосибирск", "Екатеринбург",
                            "Нижний", "Казан", "Челябинск", "Омск", "Самар",
                            "Ростов", "Уф", "Красноярск", "Перм", "Воронеж",
                            "Волгоград", "Тул", "Томск", "Новгород", "г."]
            if not any(c in str(value) for c in known_cities):
                confidence = 0.6
                parsed["warnings"].append(
                    f"Город не распознан как известный: '{value}'"
                )
        
        # Проверка направления: должно быть из известного списка
        if field == "direction_hint":
            known_directions = ["CARBIDE-STANDARD", "CARBIDE-SPECIAL", "HSS", "DIAMOND",
                                "SPEC-DRAWING", "MEASUREMENT", "TOOLING", "OTHER"]
            if str(value) not in known_directions:
                confidence = 0.6
                parsed["warnings"].append(
                    f"Направление не из стандартного списка: '{value}'"
                )
        
        # Если confidence низкая — OCR-верификация
        if confidence < 0.85 and pdf_files:
            for kw in config["ocr_keywords"]:
                for pdf_path in pdf_files:
                    ocr_text = ocr_region_around_keyword(pdf_path, kw, margin_px=250)
                    if ocr_text and len(ocr_text) > 10:
                        ocr_parsed = parse_tender_fields(ocr_text, [])
                        if ocr_parsed.get(field) is not None:
                            if str(ocr_parsed[field]) != str(parsed[field]):
                                parsed["warnings"].append(
                                    f"OCR-верификация '{config['label']}': "
                                    f"текст='{parsed[field]}' vs OCR='{ocr_parsed[field]}' — РАСХОЖДЕНИЕ"
                                )
                                confidence = 0.4
                            else:
                                confidence = 0.9  # OCR подтвердил
                            ocr_applied.append(f"{field}_verify")
                        break
                if f"{field}_verify" in ocr_applied:
                    break
        
        confidence_scores[field] = confidence
    
    # Итоговый статус
    if missing_critical:
        parsed["validation_status"] = "blocked"
    elif any(c < 0.85 for c in confidence_scores.values()):
        parsed["validation_status"] = "warnings"
    else:
        parsed["validation_status"] = "ok"
    
    parsed["confidence_scores"] = confidence_scores
    parsed["ocr_applied"] = ocr_applied
    
    return parsed


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(
        description="Быстрая экстракция + парсинг тендерных файлов (~3 сек)"
    )
    parser.add_argument("folder", nargs="?", help="Путь к папке с файлами тендера")
    parser.add_argument("--files", nargs="+", help="Список файлов (вместо папки)")
    parser.add_argument("--json", action="store_true", help="Вывести только JSON")
    parser.add_argument("--command", action="store_true",
                        help="Сгенерировать команду для emulate_llm_with_dedup.py")
    parser.add_argument("--tender-path", default="", help="Путь тендера на Яндекс.Диске")

    args = parser.parse_args()

    # Собрать список файлов
    if args.files:
        files = [os.path.abspath(f) for f in args.files if os.path.isfile(f)]
    elif args.folder:
        folder = Path(args.folder)
        if not folder.is_dir():
            print(f"Ошибка: {folder} не является директорией", file=sys.stderr)
            sys.exit(1)
        files = [str(f) for f in folder.iterdir()
                 if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS]
    else:
        print("Укажите папку или --files", file=sys.stderr)
        sys.exit(1)

    if not files:
        print("Файлы не найдены!", file=sys.stderr)
        sys.exit(1)

    # Экстракция
    total_start = time.time()
    all_text = ""
    extraction_stats = []
    file_names = [os.path.basename(f) for f in files]
    pdf_files = []  # сохраняем PDF для возможного OCR-fallback

    for filepath in files:
        text, elapsed_ms, method = extract_file(filepath)
        all_text += f"\n\n{'='*40}\n{os.path.basename(filepath)}\n{'='*40}\n{text}"
        extraction_stats.append({
            "name": os.path.basename(filepath),
            "ext": Path(filepath).suffix.lower(),
            "size_kb": os.path.getsize(filepath) / 1024,
            "time_ms": elapsed_ms,
            "chars": len(text),
            "method": method,
        })
        if Path(filepath).suffix.lower() == ".pdf":
            pdf_files.append(filepath)

    total_time = time.time() - total_start

    # Парсинг
    t0 = time.time()
    parsed = parse_tender_fields(all_text, file_names)
    parse_time = time.time() - t0

    # Валидация + OCR-fallback для ненайденных полей
    t_val = time.time()
    parsed = validate_and_ocr_fallback(parsed, pdf_files, all_text)
    validate_time = time.time() - t_val

    # Вывод
    if args.json:
        parsed["_meta"] = {
            "extraction_time_ms": round(total_time * 1000),
            "parse_time_ms": round(parse_time * 1000),
            "validate_time_ms": round(validate_time * 1000),
            "total_time_ms": round((total_time + parse_time + validate_time) * 1000),
            "files_count": len(files),
            "total_chars": sum(s["chars"] for s in extraction_stats),
            "methods": [s["method"] for s in extraction_stats],
        }
        print(json.dumps(parsed, ensure_ascii=False, indent=2))
    else:
        print(format_result(parsed, extraction_stats, total_time))
        print(f"  ⏱  Парсинг полей: {parse_time*1000:.1f} ms")
        print(f"  ⏱  Общее время: {(total_time + parse_time)*1000:.0f} ms")
        print()

        if args.command:
            print("  📝 Команда для создания сделки:")
            print("  " + "-" * 60)
            print(build_emulate_command(parsed, files, args.tender_path))
            print()


if __name__ == "__main__":
    main()
