#!/usr/bin/env python3
"""
chunk_score_extractor.py — Извлечение данных из тендерных документов
с использованием ЧАНКОВ (250 токенов) и БАЛЛЬНОЙ СИСТЕМЫ.

Архитектура:
1. Текст каждого документа разбивается на чанки (~250 токенов = ~1000 символов)
2. Каждый чанк оценивается по релевантности для каждого поля (scoring)
3. Из лучших чанков извлекаются значения полей (extraction)
4. Результаты агрегируются с учётом confidence
5. При низкой confidence — OCR fallback
"""

import logging
import os
import re
import subprocess
import time
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# ═══════════════════════════════════════════════════════════════════
# КОНФИГУРАЦИЯ ЧАНКОВ (из PATCH13 — 250 токенов)
# ═══════════════════════════════════════════════════════════════════

CHARS_PER_TOKEN = 4
CHUNK_SIZE_TOKENS = 250
CHUNK_OVERLAP_TOKENS = 75
CHUNK_SIZE = CHUNK_SIZE_TOKENS * CHARS_PER_TOKEN  # ~1000 символов
CHUNK_OVERLAP = CHUNK_OVERLAP_TOKENS * CHARS_PER_TOKEN  # ~300 символов
MIN_CHUNK_SIZE = 50  # минимальный размер чанка

# Пороги confidence
CONFIDENCE_HIGH = 0.90
CONFIDENCE_MEDIUM = 0.70
CONFIDENCE_LOW = 0.50

# Веса полей для балльной системы
FIELD_WEIGHTS = {
    "customer": {
        "keywords": ["заказчик", "организатор", "покупатель", "наименование организации",
                     "полное наименование", "АО", "ПАО", "ООО", "ФГУП", "ГК", "ФКП"],
        "context_keywords": ["ИНН", "ОГРН", "адрес", "юридический"],
        "anti_keywords": ["поставка", "инструмент", "фреза", "сверло"],
        "weight": 1.5,
    },
    "nmc": {
        "keywords": ["НМЦ", "начальная максимальная цена", "цена договора",
                     "общая стоимость", "максимальная цена", "итого с НДС", "сумма"],
        "context_keywords": ["руб", "₽", "RUB", "НДС", "без НДС"],
        "anti_keywords": [],
        "weight": 1.5,
    },
    "deadline": {
        "keywords": ["окончание подачи", "срок подачи", "дата окончания",
                     "подача заявок до", "прием заявок до", "окончание срока"],
        "context_keywords": ["заявок", "предложений", "котировочных"],
        "anti_keywords": ["поставки", "исполнения", "оплаты"],
        "weight": 1.5,
    },
    "platform": {
        "keywords": ["площадка", "ЭТП", "электронная торговая", "торговая площадка"],
        "context_keywords": ["etprf", "b2b-center", "zakupki.gov", "roseltorg",
                             "fabrikant", "rts-tender", "sberbank-ast", "tektorg"],
        "anti_keywords": [],
        "weight": 1.0,
    },
    "procedure_number": {
        "keywords": ["номер закупки", "реестровый номер", "номер процедуры",
                     "номер извещения", "регистрационный номер"],
        "context_keywords": ["закупк", "процедур", "извещени"],
        "anti_keywords": [],
        "weight": 1.0,
    },
    "procedure_type": {
        "keywords": ["способ закупки", "вид процедуры", "тип закупки",
                     "способ определения", "форма торгов"],
        "context_keywords": ["котировок", "аукцион", "конкурс", "редукцион",
                             "единственный поставщик"],
        "anti_keywords": [],
        "weight": 0.8,
    },
    "subject": {
        "keywords": ["предмет закупки", "наименование закупки", "объект закупки",
                     "наименование заказа", "предмет договора"],
        "context_keywords": ["поставка", "выполнение", "оказание"],
        "anti_keywords": [],
        "weight": 0.8,
    },
    "position_count": {
        "keywords": ["количество позиций", "кол-во лотов", "спецификация",
                     "перечень", "ведомость"],
        "context_keywords": ["№ п/п", "наименование", "ед. изм", "количество"],
        "anti_keywords": [],
        "weight": 0.7,
    },
    "equivalent_allowed": {
        "keywords": ["эквивалент", "аналог", "или эквивалент"],
        "context_keywords": ["допускается", "не допускается", "запрещен"],
        "anti_keywords": [],
        "weight": 0.7,
    },
    "gisp_required": {
        "keywords": ["ГИСП", "gisp.gov.ru", "реестр промышленной продукции"],
        "context_keywords": ["реестр", "подтверждение", "сертификат"],
        "anti_keywords": [],
        "weight": 0.7,
    },
    "city": {
        "keywords": ["место поставки", "адрес поставки", "место доставки",
                     "пункт назначения", "грузополучатель"],
        "context_keywords": ["г.", "город", "область", "край"],
        "anti_keywords": [],
        "weight": 0.6,
    },
    "delivery_terms": {
        "keywords": ["срок поставки", "срок исполнения", "срок выполнения",
                     "период поставки", "график поставки"],
        "context_keywords": ["дней", "месяц", "с момента", "после"],
        "anti_keywords": [],
        "weight": 0.5,
    },
}

DIRECTION_KEYWORDS = {
    "SPEC-DRAWING": ["по чертеж", "по тз", "по эскиз", "по образц", "нестандартн", "специнструмент"],
    "HSS-STANDARD": ["гост", "метчик", "сверло ", "развертка", "развёртка", "плашка", "зенкер",
                     "hss", "р6м5", "р18", "быстрорез"],
    "CARBIDE-STANDARD": ["твердосплав", "фреза ", "фрезы ", "пластин", "carbide", "vhm",
                         "концев", "торцев", "гравер", "promatool", "gesac", "korloy",
                         "sandvik", "iscar", "kennametal", "mitsubishi", "tungaloy"],
    "DIAMOND-STANDARD": ["алмаз", "diamond", "pcd", "cbn", "кубический нитрид бора"],
    "OUT-OF-SCOPE": ["калибр", "скоба", "штангенциркуль", "микрометр", "индикатор"],
}


# ═══════════════════════════════════════════════════════════════════
# ЧАНКИРОВАНИЕ (250 токенов, overlap 75 — из PATCH13)
# ═══════════════════════════════════════════════════════════════════

def split_into_chunks(text: str, source_file: str, file_type: str) -> list:
    """
    Разбить текст на чанки ~250 токенов с overlap 75 токенов.
    Разбивает по параграфам, сохраняя семантические границы.
    """
    if not text or len(text.strip()) < MIN_CHUNK_SIZE:
        return []

    chunks = []
    paragraphs = text.split("\n")
    current_chunk = ""
    current_start = 0
    char_pos = 0

    for para in paragraphs:
        para_len = len(para) + 1
        if len(current_chunk) + para_len > CHUNK_SIZE and current_chunk.strip():
            chunks.append({
                "text": current_chunk.strip(),
                "source_file": source_file,
                "chunk_idx": len(chunks),
                "start_char": current_start,
                "end_char": current_start + len(current_chunk),
                "file_type": file_type,
            })
            overlap_start = max(0, len(current_chunk) - CHUNK_OVERLAP)
            current_chunk = current_chunk[overlap_start:] + para + "\n"
            current_start = char_pos - (len(current_chunk) - para_len - 1)
        else:
            if not current_chunk:
                current_start = char_pos
            current_chunk += para + "\n"
        char_pos += para_len

    if current_chunk.strip() and len(current_chunk.strip()) >= MIN_CHUNK_SIZE:
        chunks.append({
            "text": current_chunk.strip(),
            "source_file": source_file,
            "chunk_idx": len(chunks),
            "start_char": current_start,
            "end_char": current_start + len(current_chunk),
            "file_type": file_type,
        })

    return chunks


# ═══════════════════════════════════════════════════════════════════
# БАЛЛЬНАЯ СИСТЕМА (SCORING)
# ═══════════════════════════════════════════════════════════════════

def score_chunk_for_field(chunk: dict, field_name: str) -> float:
    """Оценить релевантность чанка для поля (0.0 - 1.0)."""
    config = FIELD_WEIGHTS.get(field_name)
    if not config:
        return 0.0

    text_lower = chunk["text"].lower()
    score = 0.0

    keyword_hits = sum(1 for kw in config["keywords"] if kw.lower() in text_lower)
    score += min(keyword_hits * 0.3, 0.6)

    context_hits = sum(1 for kw in config["context_keywords"] if kw.lower() in text_lower)
    score += min(context_hits * 0.15, 0.3)

    anti_hits = sum(1 for kw in config.get("anti_keywords", []) if kw.lower() in text_lower)
    score -= anti_hits * 0.2

    if anti_hits == 0 and keyword_hits > 0:
        score += 0.1

    # Бонус за тип файла
    ft = chunk["file_type"].lower()
    if field_name == "nmc" and ft in ("xlsx", "xls"):
        score += 0.1
    if field_name == "customer" and ft == "pdf":
        score += 0.05
    if field_name == "position_count" and ft in ("xlsx", "xls"):
        score += 0.1

    # Бонус за позицию (первые чанки = метаданные)
    if chunk["chunk_idx"] <= 2:
        score += 0.05

    return max(0.0, min(1.0, score))


# ═══════════════════════════════════════════════════════════════════
# ИЗВЛЕЧЕНИЕ ПОЛЕЙ ИЗ ЧАНКОВ
# ═══════════════════════════════════════════════════════════════════

def extract_customer(chunk: dict) -> Optional[tuple]:
    """Извлечь заказчика. Returns (value, confidence) or None."""
    text = chunk["text"]
    patterns = [
        r'((?:Акционерное общество|Публичное акционерное общество|Общество с ограниченной ответственностью|'
        r'Федеральное государственное унитарное предприятие|Государственная корпорация)'
        r'\s*[«„"\'\u201e\u201c].*?[»"\u201d\u201f\'])',
        r'((?:АО|ПАО|ООО|ФГУП|ФКП|ГК)\s*[«„"\'\u201e\u201c][^»"\n]{3,80}[»"\u201d\u201f\'])',
        r'(?:Заказчик|Организатор|Покупатель)\s*[:\-—]\s*(.+?)(?:\n|$)',
        r'(?:Полное наименование|Наименование организации)\s*[:\-—]?\s*(.+?)(?:\n|$)',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            customer = m.group(1).strip()
            if 5 < len(customer) < 200 and not any(w in customer.lower() for w in ["поставка", "инструмент", "фреза"]):
                conf = 0.95 if re.match(r'(?:АО|ПАО|ООО|ФГУП|ГК|ФКП|Акционерное|Публичное|Общество|Федеральное)', customer) else 0.85
                return (customer, conf)
    return None


def extract_nmc(chunk: dict) -> Optional[tuple]:
    """Извлечь НМЦ. Returns (value, confidence) or None."""
    text = chunk["text"]
    patterns = [
        r'(?:НМЦ|Начальная.*?максимальная.*?цена|Цена договора|Общая стоимость)\s*[:\-—]?\s*([\d\s.,]+)\s*(?:руб|₽|RUB)',
        r'(?:Итого|ИТОГО|Всего)\s*(?:с\s*НДС)?\s*[:\-—]?\s*([\d\s.,]+)\s*(?:руб|₽)',
        r'(?:Максимальная цена)\s*[:\-—]?\s*([\d\s.,]+)',
        r'НМЦ\s*\(руб\.?\)\s*[:\-—]?\s*([\d\s.,]+)',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            nmc_str = re.sub(r'\s+', '', m.group(1)).replace(',', '.')
            parts = nmc_str.split('.')
            if len(parts) > 2:
                nmc_str = ''.join(parts[:-1]) + '.' + parts[-1]
            try:
                nmc = float(nmc_str)
                if nmc > 1000:
                    conf = 0.90 if 100_000 <= nmc <= 10_000_000_000 else 0.65
                    return (nmc, conf)
            except ValueError:
                continue
    return None


def extract_deadline(chunk: dict) -> Optional[tuple]:
    """Извлечь дедлайн. Returns (value, confidence) or None."""
    text = chunk["text"]
    patterns = [
        r'(?:Окончание.*?подачи|Срок подачи|Дата окончания подачи|Подача.*?до|Прием заявок до)\s*(?:заявок)?\s*[:\-—]?\s*(\d{1,2}[./]\d{1,2}[./]\d{2,4})\s*(\d{1,2}:\d{2})?',
        r'(?:Окончание срока подачи заявок)\s+(\d{1,2}\.\d{1,2}\.\d{4})\s*(\d{1,2}:\d{2})?',
        r'(?:до|не позднее)\s+(\d{1,2}[./]\d{1,2}[./]\d{2,4})\s*(?:г\.?)?\s*(\d{1,2}:\d{2})?',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            date_str = m.group(1).replace('/', '.')
            parts = date_str.split('.')
            if len(parts) == 3:
                d, mo, y = parts
                if len(y) == 2:
                    y = "20" + y
                try:
                    year = int(y)
                    conf = 0.90 if 2024 <= year <= 2030 else 0.50
                    return (f"{y}-{mo.zfill(2)}-{d.zfill(2)}", conf)
                except (ValueError, IndexError):
                    continue
    return None


def extract_platform(chunk: dict) -> Optional[tuple]:
    """Извлечь площадку."""
    text_lower = chunk["text"].lower()
    platforms = {
        "etprf.ru": "ЕТПРФ", "b2b-center.ru": "B2B-Center",
        "zakupki.gov.ru": "Госзакупки (ЕИС)", "roseltorg.ru": "Росэлторг",
        "etp-ets.ru": "ЕТС", "fabrikant.ru": "Фабрикант",
        "rts-tender.ru": "РТС-тендер", "sberbank-ast.ru": "Сбербанк-АСТ",
        "lot-online.ru": "Lot-Online", "tektorg.ru": "ТЭК-Торг",
        "astgoz.ru": "АСТ ГОЗ",
    }
    for domain, name in platforms.items():
        if domain in text_lower:
            return (name, 0.95)
    m = re.search(r'(?:Электронная торговая площадка|ЭТП|Площадка)\s*[:\-—]?\s*[«"]?(.+?)[»"]?\s*(?:\n|$)', chunk["text"], re.IGNORECASE)
    if m:
        return (m.group(1).strip()[:50], 0.80)
    return None


def extract_procedure_number(chunk: dict) -> Optional[tuple]:
    """Извлечь номер процедуры."""
    text = chunk["text"]
    patterns = [
        r'(?:Номер закупки|Реестровый номер|Номер процедуры|Номер извещения)\s*[:\-—]?\s*(\d[\d\-./]+\d)',
        r'(\d{4}-\d{4}-\d{5})',
        r'(?:извещени[еяю])\s*(?:№|#)?\s*(\d{10,})',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            num = m.group(1).strip()
            conf = 0.95 if re.match(r'\d{4}-\d{4}-\d{5}', num) else 0.85
            return (num, conf)
    return None


def extract_procedure_type(chunk: dict) -> Optional[tuple]:
    """Извлечь тип процедуры."""
    text = chunk["text"]
    patterns = [
        r'(?:Способ закупки|Вид процедуры|Тип закупки|Способ определения)\s*[:\-—]?\s*(.+?)(?:\n|$)',
        r'((?:Закрытый|Открытый)\s+(?:запрос\s+котировок|аукцион|конкурс|редукцион))',
        r'(Запрос\s+(?:котировок|предложений))',
    ]
    known = ["запрос котировок", "аукцион", "конкурс", "запрос предложений", "единственный поставщик", "редукцион"]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            val = m.group(1).strip()[:80]
            conf = 0.90 if any(k in val.lower() for k in known) else 0.70
            return (val, conf)
    return None


def extract_subject(chunk: dict) -> Optional[tuple]:
    """Извлечь предмет закупки."""
    text = chunk["text"]
    patterns = [
        r'Наименование заказа\s+(.+?)(?:\n|$)',
        r'(?:Предмет|Наименование закупки|Объект закупки)\s*[:\-—]?\s*(.+?)(?:\n|$)',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            subj = m.group(1).strip()
            if len(subj) > 10:
                return (subj[:200], 0.85)
    return None


def extract_position_count(chunk: dict) -> Optional[tuple]:
    """Извлечь количество позиций."""
    nums = re.findall(r'^\s*(\d{1,3})\s*[.\t|]', chunk["text"], re.MULTILINE)
    if nums:
        max_pos = max(int(n) for n in nums if 1 <= int(n) <= 500)
        if max_pos >= 2:
            return (max_pos, 0.85 if max_pos >= 5 else 0.75)
    return None


def extract_equivalent(chunk: dict) -> Optional[tuple]:
    """Извлечь информацию об эквиваленте."""
    text = chunk["text"]
    if re.search(r'(?:или\s+эквивалент|эквивалент\s*допускается|аналог\s*допускается)', text, re.IGNORECASE):
        return (True, 0.85)
    if re.search(r'(?:без\s*эквивалент|эквивалент\s*не\s*допускается)', text, re.IGNORECASE):
        return (False, 0.85)
    return None


def extract_gisp(chunk: dict) -> Optional[tuple]:
    """Извлечь требование ГИСП."""
    if re.search(r'(?:ГИСП|gisp\.gov\.ru|реестр\s*промышленной\s*продукции)', chunk["text"], re.IGNORECASE):
        return (True, 0.90)
    return None


def extract_city(chunk: dict) -> Optional[tuple]:
    """Извлечь город."""
    m = re.search(
        r'(?:г\.\s*|город\s+)(Москв[аы]|Санкт-Петербург[а]?|Нижн[а-я]*\s*Новгород[а]?|'
        r'Екатеринбург[а]?|Новосибирск[а]?|Казан[ьи]|Челябинск[а]?|Омск[а]?|'
        r'Самар[аы]|Ростов[а]?(?:\s*-на-Дону)?|Уф[аы]|Красноярск[а]?|'
        r'Пермь|Перми|Воронеж[а]?|Волгоград[а]?|Тул[аы]|Ижевск[а]?|'
        r'Курган[а]?|Ковров[а]?|Томск[а]?|[А-ЯЁ][а-яё]+)',
        chunk["text"]
    )
    if m and len(m.group(1)) > 2:
        return (f"г. {m.group(1)}", 0.80)
    return None


def extract_delivery_terms(chunk: dict) -> Optional[tuple]:
    """Извлечь сроки поставки."""
    patterns = [
        r'(?:Срок поставки|Срок исполнения)\s*[:\-—]?\s*(.+?)(?:\n|$)',
        r'(\d+)\s*(?:календарных|рабочих)?\s*дн[а-я]*\s*(?:с момента|после|от даты)',
    ]
    for pat in patterns:
        m = re.search(pat, chunk["text"], re.IGNORECASE)
        if m:
            val = m.group(0).strip()[:100]
            if not val.endswith(('.pdf', '.docx', '.xlsx')):
                return (val, 0.75)
    return None


# Маппинг поле → функция
FIELD_EXTRACTORS = {
    "customer": extract_customer,
    "nmc": extract_nmc,
    "deadline": extract_deadline,
    "platform": extract_platform,
    "procedure_number": extract_procedure_number,
    "procedure_type": extract_procedure_type,
    "subject": extract_subject,
    "position_count": extract_position_count,
    "equivalent_allowed": extract_equivalent,
    "gisp_required": extract_gisp,
    "city": extract_city,
    "delivery_terms": extract_delivery_terms,
}


# ═══════════════════════════════════════════════════════════════════
# OCR FALLBACK
# ═══════════════════════════════════════════════════════════════════

def ocr_pdf_full(filepath: str) -> str:
    """Полный OCR для PDF."""
    try:
        from pdf2image import convert_from_path
        import pytesseract
        logger.info(f"OCR fallback: {os.path.basename(filepath)}")
        images = convert_from_path(filepath, dpi=300)
        parts = []
        for img in images:
            parts.append(pytesseract.image_to_string(img, lang='rus+eng'))
        return "\n\n".join(parts)
    except Exception as e:
        logger.warning(f"OCR ошибка: {e}")
        return ""


def ocr_region_around_keyword(filepath: str, keyword: str, margin_px: int = 250) -> str:
    """Точечный OCR вокруг ключевого слова."""
    try:
        from pdf2image import convert_from_path
        import pytesseract

        result = subprocess.run(
            ["pdftotext", "-bbox", filepath, "-"],
            capture_output=True, text=True, timeout=30
        )
        keyword_lower = keyword.lower()
        bbox_match = None
        page_num = 0
        for line in result.stdout.splitlines():
            if '<page' in line:
                page_num += 1
            if keyword_lower in line.lower() and 'xMin' in line:
                coords = re.findall(r'[xy](?:Min|Max)="([\d.]+)"', line)
                if len(coords) >= 4:
                    bbox_match = {'page': page_num, 'xMin': float(coords[0]),
                                  'yMin': float(coords[1]), 'xMax': float(coords[2]),
                                  'yMax': float(coords[3])}
                    break
        if not bbox_match:
            return ""
        images = convert_from_path(filepath, dpi=300,
                                   first_page=bbox_match['page'], last_page=bbox_match['page'])
        if not images:
            return ""
        img = images[0]
        w, h = img.size
        scale = 300 / 72
        x1 = max(0, int(bbox_match['xMin'] * scale) - margin_px)
        y1 = max(0, int(bbox_match['yMin'] * scale) - margin_px)
        x2 = min(w, int(bbox_match['xMax'] * scale) + margin_px)
        y2 = min(h, int(bbox_match['yMax'] * scale) + margin_px)
        region = img.crop((x1, y1, x2, y2))
        return pytesseract.image_to_string(region, lang='rus+eng').strip()
    except Exception as e:
        logger.warning(f"OCR region failed for '{keyword}': {e}")
        return ""


def apply_ocr_fallback(field_name: str, pdf_files: list, all_chunks: list) -> Optional[tuple]:
    """Применить OCR fallback для поля с низкой confidence."""
    config = FIELD_WEIGHTS.get(field_name)
    extractor = FIELD_EXTRACTORS.get(field_name)
    if not config or not extractor:
        return None

    # Попытка 1: точечный OCR
    for kw in config["keywords"][:3]:
        for pdf_path in pdf_files:
            ocr_text = ocr_region_around_keyword(pdf_path, kw, margin_px=300)
            if ocr_text and len(ocr_text) > 20:
                ocr_chunk = {"text": ocr_text, "source_file": pdf_path,
                             "chunk_idx": 9999, "start_char": 0,
                             "end_char": len(ocr_text), "file_type": "ocr"}
                result = extractor(ocr_chunk)
                if result:
                    return (result[0], result[1] * 0.85)  # Снижаем confidence для OCR

    # Попытка 2: полный OCR первого PDF
    if pdf_files:
        ocr_text = ocr_pdf_full(pdf_files[0])
        if ocr_text and len(ocr_text) > 100:
            ocr_chunks = split_into_chunks(ocr_text, pdf_files[0], "ocr")
            for chunk in ocr_chunks:
                result = extractor(chunk)
                if result:
                    return (result[0], result[1] * 0.75)
    return None


# ═══════════════════════════════════════════════════════════════════
# НАПРАВЛЕНИЕ И ПРИОРИТЕТ
# ═══════════════════════════════════════════════════════════════════

def score_direction(all_chunks: list) -> tuple:
    """Определить направление по балльной системе."""
    scores = {d: 0.0 for d in DIRECTION_KEYWORDS}
    for chunk in all_chunks:
        text_lower = chunk["text"].lower()
        for direction, keywords in DIRECTION_KEYWORDS.items():
            for kw in keywords:
                hits = len(re.findall(re.escape(kw), text_lower))
                if hits > 0:
                    scores[direction] += hits / max(len(chunk["text"]) / 1000, 1)

    max_dir = max(scores, key=scores.get)
    max_score = scores[max_dir]
    if max_score == 0:
        return ("CARBIDE-STANDARD", 0.30, scores)

    sorted_dirs = sorted(scores.items(), key=lambda x: -x[1])
    if len(sorted_dirs) > 1 and sorted_dirs[1][1] > 0:
        ratio = sorted_dirs[1][1] / max(sorted_dirs[0][1], 1)
        conf = 0.60 if ratio > 0.7 else (0.75 if ratio > 0.4 else 0.90)
    else:
        conf = 0.90
    return (max_dir, conf, scores)


def determine_priority(nmc: float, deadline: str) -> tuple:
    """Определить приоритет."""
    warnings = []
    if nmc > 20_000_000:
        priority = "Р2"
    elif nmc > 5_000_000:
        priority = "Р2"
    elif nmc > 1_000_000:
        priority = "Р3"
    else:
        priority = "Р3"

    if deadline:
        from datetime import datetime
        try:
            dl = datetime.strptime(deadline, "%Y-%m-%d")
            hours_left = (dl - datetime.now()).total_seconds() / 3600
            if hours_left <= 48:
                priority = "Р1"
                warnings.append(f"СРОЧНО: до дедлайна {hours_left:.0f}ч")
            elif hours_left <= 120:
                warnings.append(f"Внимание: до дедлайна {hours_left/24:.0f} дней")
        except ValueError:
            pass
    return (priority, warnings)


# ═══════════════════════════════════════════════════════════════════
# ИЗВЛЕЧЕНИЕ ТЕКСТА ИЗ ФАЙЛОВ
# ═══════════════════════════════════════════════════════════════════

def extract_file_text(filepath: str) -> str:
    """Извлечь текст из файла."""
    ext = Path(filepath).suffix.lower()
    try:
        if ext == ".pdf":
            result = subprocess.run(["pdftotext", "-layout", filepath, "-"],
                                    capture_output=True, text=True, timeout=30)
            text = result.stdout
            if len(text.strip()) < 200:
                return ocr_pdf_full(filepath)
            return text
        elif ext == ".docx":
            from docx import Document
            doc = Document(filepath)
            text = "\n".join([p.text for p in doc.paragraphs])
            for table in doc.tables:
                for row in table.rows:
                    text += "\n" + "\t".join([cell.text for cell in row.cells])
            return text
        elif ext in (".xlsx", ".xls"):
            import openpyxl
            wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
            text = ""
            for sheet in wb.sheetnames:
                ws = wb[sheet]
                text += f"\n=== {sheet} ===\n"
                for row in ws.iter_rows(values_only=True):
                    text += "\t".join(str(v) if v else "" for v in row) + "\n"
            wb.close()
            return text
        elif ext in (".txt", ".csv"):
            return Path(filepath).read_text(encoding="utf-8", errors="replace")
        elif ext == ".doc":
            result = subprocess.run(["antiword", filepath],
                                    capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return result.stdout
            result = subprocess.run(["catdoc", filepath],
                                    capture_output=True, text=True, timeout=30)
            return result.stdout
    except Exception as e:
        logger.warning(f"Ошибка извлечения {filepath}: {e}")
    return ""


# ═══════════════════════════════════════════════════════════════════
# ГЛАВНАЯ ФУНКЦИЯ
# ═══════════════════════════════════════════════════════════════════

def extract_and_parse_tender(all_files: list) -> dict:
    """
    Главная функция: извлечение данных из тендерных документов
    с чанками (250 токенов) и балльной системой.

    Args:
        all_files: список абсолютных путей к скачанным файлам

    Returns:
        dict совместимый с cron_yadisk._create_amo_lead()
    """
    t0 = time.time()

    # 1. Извлечение текста и чанкирование
    all_chunks = []
    file_names = []
    pdf_files = []

    for filepath in all_files:
        if not os.path.exists(filepath):
            continue
        filename = os.path.basename(filepath)
        file_names.append(filename)
        ext = Path(filepath).suffix.lower()
        if ext == ".pdf":
            pdf_files.append(filepath)

        text = extract_file_text(filepath)
        if text and len(text.strip()) > MIN_CHUNK_SIZE:
            file_type = ext.lstrip('.')
            chunks = split_into_chunks(text, filename, file_type)
            all_chunks.extend(chunks)

    total_chunks = len(all_chunks)
    if not all_chunks:
        return {
            "customer": None, "nmc": 0, "direction": "CARBIDE-STANDARD",
            "deadline": "", "priority": "P3", "validation_status": "blocked",
            "confidence": {}, "warnings": ["Не удалось извлечь текст"],
            "chunks_used": 0, "total_chunks": 0,
        }

    logger.info(f"Создано {total_chunks} чанков из {len(file_names)} файлов")

    # 2. Извлечение полей с scoring
    fields = {}
    confidence_scores = {}
    chunks_used = set()
    ocr_applied = []
    warnings = []

    for field_name, extractor in FIELD_EXTRACTORS.items():
        # Score all chunks for this field
        scored = [(chunk, score_chunk_for_field(chunk, field_name))
                  for chunk in all_chunks]
        scored = [(c, s) for c, s in scored if s > 0.1]
        scored.sort(key=lambda x: -x[1])

        # Extract from top-5 chunks
        best_value = None
        best_conf = 0.0
        for chunk, chunk_score in scored[:5]:
            result = extractor(chunk)
            if result:
                value, conf = result
                final_conf = conf * (0.5 + chunk_score * 0.5)
                if final_conf > best_conf:
                    best_value = value
                    best_conf = final_conf
                    chunks_used.add(chunk["chunk_idx"])

        fields[field_name] = best_value
        confidence_scores[field_name] = best_conf

    # 3. OCR fallback для критичных полей
    critical_fields = ["customer", "nmc", "deadline"]
    for field_name in critical_fields:
        if (fields.get(field_name) is None or confidence_scores.get(field_name, 0) < CONFIDENCE_LOW) and pdf_files:
            logger.info(f"OCR fallback для '{field_name}' (conf={confidence_scores.get(field_name, 0):.2f})")
            ocr_result = apply_ocr_fallback(field_name, pdf_files, all_chunks)
            if ocr_result:
                value, conf = ocr_result
                if conf > confidence_scores.get(field_name, 0):
                    fields[field_name] = value
                    confidence_scores[field_name] = conf
                    ocr_applied.append(field_name)
                    warnings.append(f"Поле '{field_name}' извлечено через OCR")

    # 4. Направление
    direction, dir_conf, dir_scores = score_direction(all_chunks)

    # 5. Приоритет
    nmc_val = fields.get("nmc") or 0
    if isinstance(nmc_val, str):
        try:
            nmc_val = float(nmc_val)
        except ValueError:
            nmc_val = 0
    priority, prio_warnings = determine_priority(nmc_val, fields.get("deadline") or "")
    warnings.extend(prio_warnings)

    # 6. Валидация
    missing = [f for f in critical_fields if not fields.get(f)]
    if missing:
        validation_status = "blocked"
        warnings.append(f"КРИТИЧНО: не найдены: {', '.join(missing)}")
    elif any(confidence_scores.get(f, 0) < CONFIDENCE_MEDIUM for f in critical_fields):
        validation_status = "warnings"
    else:
        validation_status = "ok"

    # 7. Доп. обработка
    customer_short = ""
    if fields.get("customer"):
        short = re.sub(r'^(Акционерное общество|Публичное акционерное общество|'
                       r'Общество с ограниченной ответственностью|'
                       r'Федеральное государственное унитарное предприятие|'
                       r'Государственная корпорация|АО|ПАО|ООО|ФГУП|ФКП|ГК)\s*',
                       '', fields["customer"], flags=re.IGNORECASE)
        customer_short = re.sub(r'[«»„""\u201e\u201c\u201d]+', '', short).strip()[:50]

    nmc_formatted = f"{nmc_val:,.2f} руб.".replace(',', ' ') if nmc_val > 0 else ""

    # Ключевые продукты
    products = set()
    for chunk in all_chunks:
        for kw in re.findall(r'(?:фреза|фрезы|метчик|сверло|развертка|развёртка|пластина|гравер|'
                             r'зенкер|плашка|резец|борфреза|державка|патрон|оправка)\w*',
                             chunk["text"].lower()):
            products.add(kw[:20])

    elapsed = (time.time() - t0) * 1000

    return {
        "customer": fields.get("customer"),
        "customer_short": customer_short,
        "nmc": nmc_val,
        "nmc_formatted": nmc_formatted,
        "direction": direction,
        "deadline": fields.get("deadline") or "",
        "priority": priority,
        "validation_status": validation_status,
        "confidence": confidence_scores,
        "warnings": warnings,
        "ocr_applied": ocr_applied,
        "chunks_used": len(chunks_used),
        "total_chunks": total_chunks,
        "extraction_time_ms": elapsed,
        "platform": fields.get("platform"),
        "procedure_number": fields.get("procedure_number"),
        "procedure_type": fields.get("procedure_type"),
        "subject": fields.get("subject"),
        "position_count": fields.get("position_count"),
        "equivalent_allowed": fields.get("equivalent_allowed"),
        "gisp_required": fields.get("gisp_required"),
        "city": fields.get("city"),
        "delivery_terms": fields.get("delivery_terms"),
        "key_products": sorted(list(products))[:10],
    }


# ═══════════════════════════════════════════════════════════════════
# CLI
# ═══════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import json
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

    if len(sys.argv) < 2:
        print("Usage: python3 chunk_score_extractor.py /path/to/folder")
        print("       python3 chunk_score_extractor.py file1.pdf file2.docx")
        sys.exit(1)

    path = Path(sys.argv[1])
    if path.is_dir():
        files = []
        for ext in [".pdf", ".docx", ".xlsx", ".xls", ".doc", ".txt", ".csv"]:
            files.extend(str(f) for f in path.glob(f"*{ext}"))
    else:
        files = sys.argv[1:]

    result = extract_and_parse_tender(files)

    print("=" * 70)
    print("  РЕЗУЛЬТАТ (ЧАНКИ 250 токенов + БАЛЛЬНАЯ СИСТЕМА)")
    print("=" * 70)
    print(f"  Заказчик:     {result.get('customer') or '❌ НЕ НАЙДЕН'}")
    print(f"  НМЦ:          {result.get('nmc_formatted') or '❌'}")
    print(f"  Дедлайн:      {result.get('deadline') or '❌'}")
    print(f"  Площадка:     {result.get('platform') or '—'}")
    print(f"  Номер:        {result.get('procedure_number') or '—'}")
    print(f"  Тип:          {result.get('procedure_type') or '—'}")
    print(f"  Направление:  {result.get('direction')}")
    print(f"  Приоритет:    {result.get('priority')}")
    print(f"  Статус:       {result.get('validation_status')}")
    print(f"  Чанков:       {result.get('chunks_used')}/{result.get('total_chunks')}")
    print(f"  OCR:          {result.get('ocr_applied') or 'не потребовался'}")
    print(f"  Время:        {result.get('extraction_time_ms', 0):.0f} ms")
    if result.get('warnings'):
        print(f"\n  ⚠️ Предупреждения:")
        for w in result['warnings']:
            print(f"    - {w}")
    print(f"\n  Confidence:")
    for field, score in sorted(result.get('confidence', {}).items(), key=lambda x: -x[1]):
        bar = "█" * int(score * 10) + "░" * (10 - int(score * 10))
        print(f"    {field:20s} [{bar}] {score:.2f}")
    print("=" * 70)
    print("\nJSON:")
    print(json.dumps(result, ensure_ascii=False, indent=2, default=str))
