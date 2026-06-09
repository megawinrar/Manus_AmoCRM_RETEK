"""
Бенчмарк: замер времени извлечения данных из файлов тендера.
Цель — понять, где узкое место и нужна ли оптимизация.
"""

import time
import os
import subprocess

UPLOAD_DIR = "/home/ubuntu/upload"

FILES = [
    "2109-2026-00743.Поставкаинструментапроизводства_Promatool_илиэквивалент—Закупкавх.№195871.pdf",
    "2026-06-08_12-50-12_ИДоЗ.docx",
    "Проектдоговора№0730_8от29.04.2026(41305696v2).docx",
    "СведенияоНМЦзакупки№0730_6от20.04.2026(40698748v6).xlsx",
    "Техническоезаданиеназакупку№0730_8от21.04.2026(40782358v2).docx",
    "Форма046_ЗЗК.docx",
]


def extract_pdf(filepath):
    """Извлечь текст из PDF через pdftotext (poppler)."""
    result = subprocess.run(
        ["pdftotext", "-layout", filepath, "-"],
        capture_output=True, text=True, timeout=30
    )
    return result.stdout


def extract_docx(filepath):
    """Извлечь текст из DOCX через python-docx."""
    from docx import Document
    doc = Document(filepath)
    text = "\n".join([p.text for p in doc.paragraphs])
    # Также извлечь таблицы
    for table in doc.tables:
        for row in table.rows:
            text += "\n" + "\t".join([cell.text for cell in row.cells])
    return text


def extract_xlsx(filepath):
    """Извлечь данные из XLSX через openpyxl."""
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True, read_only=True)
    text = ""
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text += f"\n=== Sheet: {sheet_name} ===\n"
        for row in ws.iter_rows(values_only=True):
            vals = [str(v) if v is not None else "" for v in row]
            text += "\t".join(vals) + "\n"
    wb.close()
    return text


def main():
    print("=" * 60)
    print("БЕНЧМАРК ИЗВЛЕЧЕНИЯ ДАННЫХ ИЗ ФАЙЛОВ ТЕНДЕРА")
    print("=" * 60)
    print()

    total_start = time.time()
    results = []

    for fname in FILES:
        fpath = os.path.join(UPLOAD_DIR, fname)
        fsize = os.path.getsize(fpath)
        ext = os.path.splitext(fname)[1].lower()

        t0 = time.time()

        if ext == ".pdf":
            text = extract_pdf(fpath)
        elif ext == ".docx":
            text = extract_docx(fpath)
        elif ext == ".xlsx":
            text = extract_xlsx(fpath)
        else:
            text = ""

        elapsed = time.time() - t0
        chars = len(text)

        results.append({
            "file": fname[:50],
            "size_kb": fsize / 1024,
            "ext": ext,
            "time_ms": elapsed * 1000,
            "chars": chars,
        })

        print(f"  {ext:6s} | {elapsed*1000:7.1f} ms | {fsize/1024:8.1f} KB | {chars:7,} chars | {fname[:55]}")

    total_elapsed = time.time() - total_start

    print()
    print("-" * 60)
    print(f"  ИТОГО: {total_elapsed*1000:.0f} ms ({total_elapsed:.2f} сек) на {len(FILES)} файлов")
    print(f"  Суммарный текст: {sum(r['chars'] for r in results):,} символов")
    print(f"  Суммарный размер: {sum(r['size_kb'] for r in results):.0f} KB")
    print("-" * 60)

    # Breakdown by type
    print()
    print("  По типам:")
    for ext in [".pdf", ".docx", ".xlsx"]:
        items = [r for r in results if r["ext"] == ext]
        if items:
            total_ms = sum(r["time_ms"] for r in items)
            print(f"    {ext}: {len(items)} файлов, {total_ms:.0f} ms суммарно")


if __name__ == "__main__":
    main()
