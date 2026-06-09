#!/usr/bin/env python3
"""
RAG Indexer — Чанкование документов + индексация в ChromaDB через Yandex Embeddings.

Архитектура:
    Файлы тендера (PDF/DOCX/XLSX)
        → extract text (pdftotext / python-docx / openpyxl)
        → chunk (500 токенов, overlap 50)
        → embed (Yandex text-search-doc)
        → store (ChromaDB, persistent)

Использование:
    # Индексировать один тендер
    python3 scripts/rag_indexer.py --files file1.pdf file2.docx --tender-id "2109-2026-00743"

    # Индексировать папку
    python3 scripts/rag_indexer.py --folder /path/to/tender --tender-id "2109-2026-00743"

    # Показать статистику базы
    python3 scripts/rag_indexer.py --stats
"""

import os
import sys
import json
import time
import asyncio
import hashlib
import argparse
import subprocess
from pathlib import Path
from typing import List, Dict, Tuple, Optional

import requests

# ═══════════════════════════════════════════════════════════════════════════
# CONFIGURATION
# ═══════════════════════════════════════════════════════════════════════════

# Load .env
ENV_PATH = Path(__file__).parent.parent / ".env"
if ENV_PATH.exists():
    for line in ENV_PATH.read_text().splitlines():
        line = line.strip()
        if line and not line.startswith("#") and "=" in line:
            key, _, val = line.partition("=")
            os.environ.setdefault(key.strip(), val.strip())

YANDEX_API_KEY = os.environ.get("YANDEX_GPT_API_KEY", "")
YANDEX_FOLDER_ID = os.environ.get("YANDEX_GPT_FOLDER_ID", "")
EMBEDDING_MODEL = os.environ.get("YANDEX_EMBEDDING_MODEL", "text-search-doc/latest")

CHROMA_DB_PATH = Path(__file__).parent.parent / "data" / "chroma_db"
CHUNK_SIZE = 250       # tokens (~1000 chars for Russian) — smaller chunks = better precision
CHUNK_OVERLAP = 75     # tokens overlap between chunks (30% overlap for context)
CHARS_PER_TOKEN = 4    # approximate for Russian text
BATCH_SIZE = 20        # concurrent embeddings per wave
RATE_LIMIT_DELAY = 0.1 # seconds between waves (not between individual calls)
MAX_CONCURRENT = 20    # max parallel requests

EMBEDDING_URL = "https://llm.api.cloud.yandex.net/foundationModels/v1/textEmbedding"


# ═══════════════════════════════════════════════════════════════════════════
# TEXT EXTRACTION (reuse from extract_and_classify.py)
# ═══════════════════════════════════════════════════════════════════════════

def extract_pdf(filepath: str) -> str:
    """Extract text from PDF using pdftotext."""
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", filepath, "-"],
            capture_output=True, text=True, timeout=30
        )
        return result.stdout
    except Exception as e:
        print(f"  [WARN] PDF extraction failed: {e}")
        return ""


def extract_docx(filepath: str) -> str:
    """Extract text from DOCX using python-docx."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        
        return "\n".join(paragraphs)
    except Exception as e:
        print(f"  [WARN] DOCX extraction failed: {e}")
        return ""


def extract_xlsx(filepath: str) -> str:
    """Extract text from XLSX using openpyxl."""
    try:
        from openpyxl import load_workbook
        wb = load_workbook(filepath, data_only=True, read_only=True)
        lines = []
        for ws in wb.worksheets:
            lines.append(f"=== Лист: {ws.title} ===")
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else "" for c in row]
                row_text = " | ".join(c for c in cells if c)
                if row_text.strip():
                    lines.append(row_text)
        wb.close()
        return "\n".join(lines)
    except Exception as e:
        print(f"  [WARN] XLSX extraction failed: {e}")
        return ""


def extract_file(filepath: str) -> str:
    """Extract text from any supported file."""
    ext = Path(filepath).suffix.lower()
    if ext == ".pdf":
        return extract_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_docx(filepath)
    elif ext in (".xlsx", ".xls"):
        return extract_xlsx(filepath)
    elif ext in (".txt", ".md", ".csv"):
        return Path(filepath).read_text(encoding="utf-8", errors="ignore")
    else:
        print(f"  [WARN] Unsupported file type: {ext}")
        return ""


# ═══════════════════════════════════════════════════════════════════════════
# CHUNKING
# ═══════════════════════════════════════════════════════════════════════════

def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> List[Dict]:
    """
    Split text into overlapping chunks.
    
    Strategy:
    - Split by paragraphs first (preserve semantic boundaries)
    - If paragraph > chunk_size, split by sentences
    - Add overlap between chunks for context continuity
    
    Returns list of dicts: {text, start_char, end_char, chunk_idx}
    """
    if not text.strip():
        return []
    
    char_chunk_size = chunk_size * CHARS_PER_TOKEN
    char_overlap = overlap * CHARS_PER_TOKEN
    
    chunks = []
    
    # Split into paragraphs
    paragraphs = text.split("\n")
    
    current_chunk = ""
    current_start = 0
    char_pos = 0
    
    for para in paragraphs:
        para_len = len(para) + 1  # +1 for \n
        
        # If adding this paragraph exceeds chunk size
        if len(current_chunk) + para_len > char_chunk_size and current_chunk.strip():
            # Save current chunk
            chunks.append({
                "text": current_chunk.strip(),
                "start_char": current_start,
                "end_char": current_start + len(current_chunk),
                "chunk_idx": len(chunks)
            })
            
            # Start new chunk with overlap
            overlap_start = max(0, len(current_chunk) - char_overlap)
            current_chunk = current_chunk[overlap_start:] + para + "\n"
            current_start = char_pos - (len(current_chunk) - para_len - 1)
        else:
            if not current_chunk:
                current_start = char_pos
            current_chunk += para + "\n"
        
        char_pos += para_len
    
    # Don't forget the last chunk
    if current_chunk.strip():
        chunks.append({
            "text": current_chunk.strip(),
            "start_char": current_start,
            "end_char": current_start + len(current_chunk),
            "chunk_idx": len(chunks)
        })
    
    return chunks


# ═══════════════════════════════════════════════════════════════════════════
# YANDEX EMBEDDINGS
# ═══════════════════════════════════════════════════════════════════════════

def get_embedding_sync(text: str, model_uri: Optional[str] = None) -> Optional[List[float]]:
    """Get embedding vector from Yandex API for a single text (sync fallback)."""
    if not model_uri:
        model_uri = f"emb://{YANDEX_FOLDER_ID}/{EMBEDDING_MODEL}"
    
    # Truncate to ~8000 tokens (~32000 chars) — API limit
    if len(text) > 32000:
        text = text[:32000]
    
    try:
        resp = requests.post(
            EMBEDDING_URL,
            headers={
                "Content-Type": "application/json",
                "Authorization": f"Api-Key {YANDEX_API_KEY}",
                "x-folder-id": YANDEX_FOLDER_ID,
                "x-data-logging-enabled": "false",
            },
            json={
                "modelUri": model_uri,
                "text": text,
            },
            timeout=30,
        )
        
        if resp.status_code == 200:
            return resp.json()["embedding"]
        elif resp.status_code == 429:
            # Rate limited — wait and retry
            time.sleep(2)
            return get_embedding_sync(text, model_uri)
        else:
            print(f"  [ERROR] Embedding API {resp.status_code}: {resp.text[:200]}")
            return None
    except Exception as e:
        print(f"  [ERROR] Embedding request failed: {e}")
        return None


async def get_embedding_async(session, text: str, semaphore, model_uri: str) -> Optional[List[float]]:
    """Get embedding vector from Yandex API (async with concurrency control)."""
    import aiohttp
    
    if len(text) > 32000:
        text = text[:32000]
    
    async with semaphore:
        try:
            async with session.post(
                EMBEDDING_URL,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Api-Key {YANDEX_API_KEY}",
                    "x-folder-id": YANDEX_FOLDER_ID,
                    "x-data-logging-enabled": "false",
                },
                json={
                    "modelUri": model_uri,
                    "text": text,
                },
                timeout=aiohttp.ClientTimeout(total=30),
            ) as resp:
                if resp.status == 200:
                    data = await resp.json()
                    return data["embedding"]
                elif resp.status == 429:
                    # Rate limited — back off and retry
                    await asyncio.sleep(2)
                    async with session.post(
                        EMBEDDING_URL,
                        headers={
                            "Content-Type": "application/json",
                            "Authorization": f"Api-Key {YANDEX_API_KEY}",
                            "x-folder-id": YANDEX_FOLDER_ID,
                            "x-data-logging-enabled": "false",
                        },
                        json={"modelUri": model_uri, "text": text},
                        timeout=aiohttp.ClientTimeout(total=30),
                    ) as retry_resp:
                        if retry_resp.status == 200:
                            data = await retry_resp.json()
                            return data["embedding"]
                        else:
                            text_body = await retry_resp.text()
                            print(f"  [ERROR] Retry failed {retry_resp.status}: {text_body[:100]}")
                            return None
                else:
                    text_body = await resp.text()
                    print(f"  [ERROR] Embedding API {resp.status}: {text_body[:200]}")
                    return None
        except Exception as e:
            print(f"  [ERROR] Embedding request failed: {e}")
            return None


async def get_embeddings_concurrent(texts: List[str]) -> List[Optional[List[float]]]:
    """Get embeddings for all texts concurrently (up to MAX_CONCURRENT parallel)."""
    import aiohttp
    
    model_uri = f"emb://{YANDEX_FOLDER_ID}/{EMBEDDING_MODEL}"
    semaphore = asyncio.Semaphore(MAX_CONCURRENT)
    
    async with aiohttp.ClientSession() as session:
        tasks = [
            get_embedding_async(session, text, semaphore, model_uri)
            for text in texts
        ]
        results = await asyncio.gather(*tasks)
    
    return list(results)


def get_embeddings_batch(texts: List[str]) -> List[Optional[List[float]]]:
    """Get embeddings — uses async concurrent if available, else sequential."""
    try:
        import aiohttp  # noqa: F401
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            results = loop.run_until_complete(get_embeddings_concurrent(texts))
        finally:
            loop.close()
        return results
    except ImportError:
        # Fallback to sequential
        print("  [INFO] aiohttp not installed, using sequential embedding")
        results = []
        for i, text in enumerate(texts):
            embedding = get_embedding_sync(text)
            results.append(embedding)
            if i < len(texts) - 1:
                time.sleep(RATE_LIMIT_DELAY)
        return results


# ═══════════════════════════════════════════════════════════════════════════
# CHROMADB STORAGE
# ═══════════════════════════════════════════════════════════════════════════

def get_collection(collection_name: str = "tenders"):
    """Get or create ChromaDB collection."""
    import chromadb
    
    CHROMA_DB_PATH.mkdir(parents=True, exist_ok=True)
    client = chromadb.PersistentClient(path=str(CHROMA_DB_PATH))
    
    collection = client.get_or_create_collection(
        name=collection_name,
        metadata={"hnsw:space": "cosine"}  # cosine similarity
    )
    
    return client, collection


def index_tender(
    files: List[str],
    tender_id: str,
    tender_name: str = "",
    metadata_extra: Optional[Dict] = None,
    collection_name: str = "tenders",
) -> Dict:
    """
    Index all files of a tender into ChromaDB.
    
    Returns stats dict with timing and counts.
    """
    stats = {
        "tender_id": tender_id,
        "files_processed": 0,
        "chunks_total": 0,
        "chunks_embedded": 0,
        "chunks_failed": 0,
        "time_extract_sec": 0,
        "time_chunk_sec": 0,
        "time_embed_sec": 0,
        "time_total_sec": 0,
        "cost_rub": 0,
    }
    
    t_start = time.time()
    
    # 1. Extract text from all files
    print(f"\n{'='*60}")
    print(f"RAG INDEXER — Тендер: {tender_id}")
    print(f"{'='*60}")
    
    t_extract = time.time()
    all_chunks = []
    
    for filepath in files:
        if not os.path.exists(filepath):
            print(f"  [SKIP] File not found: {filepath}")
            continue
        
        filename = Path(filepath).name
        print(f"  📄 {filename}...", end=" ", flush=True)
        
        text = extract_file(filepath)
        if not text.strip():
            print("EMPTY")
            continue
        
        # Chunk the text
        chunks = chunk_text(text)
        
        # Add metadata to each chunk
        for chunk in chunks:
            chunk["filename"] = filename
            chunk["filepath"] = filepath
            chunk["tender_id"] = tender_id
            chunk["tender_name"] = tender_name
            if metadata_extra:
                chunk.update(metadata_extra)
            
            # Generate unique ID
            content_hash = hashlib.md5(
                f"{tender_id}:{filename}:{chunk['chunk_idx']}:{chunk['text'][:100]}".encode()
            ).hexdigest()
            chunk["id"] = f"{tender_id}__{filename}__{chunk['chunk_idx']}__{content_hash[:8]}"
        
        all_chunks.extend(chunks)
        stats["files_processed"] += 1
        print(f"{len(chunks)} chunks, {len(text)} chars")
    
    stats["time_extract_sec"] = round(time.time() - t_extract, 2)
    stats["chunks_total"] = len(all_chunks)
    
    if not all_chunks:
        print("\n  [ERROR] No chunks to index!")
        return stats
    
    # 2. Get embeddings
    print(f"\n  🔢 Embedding {len(all_chunks)} chunks (concurrent, max {MAX_CONCURRENT} parallel)...")
    t_embed = time.time()
    
    total_tokens = 0
    embedded_chunks = []
    
    # Process in waves for progress reporting
    all_texts = [c["text"] for c in all_chunks]
    
    for wave_start in range(0, len(all_texts), BATCH_SIZE):
        wave_end = min(wave_start + BATCH_SIZE, len(all_texts))
        wave_texts = all_texts[wave_start:wave_end]
        
        embeddings = get_embeddings_batch(wave_texts)
        
        for chunk, embedding in zip(all_chunks[wave_start:wave_end], embeddings):
            if embedding is not None:
                chunk["embedding"] = embedding
                embedded_chunks.append(chunk)
                stats["chunks_embedded"] += 1
                total_tokens += len(chunk["text"]) // CHARS_PER_TOKEN
            else:
                stats["chunks_failed"] += 1
        
        # Progress
        print(f"    [{wave_end}/{len(all_chunks)}] embedded", end="\r", flush=True)
        
        # Small delay between waves
        if wave_end < len(all_texts):
            time.sleep(RATE_LIMIT_DELAY)
    
    print(f"    [{stats['chunks_embedded']}/{len(all_chunks)}] embedded — DONE")
    
    stats["time_embed_sec"] = round(time.time() - t_embed, 2)
    stats["cost_rub"] = round(total_tokens / 1000 * 0.01, 4)  # 0.01₽ per 1K tokens
    
    # 3. Store in ChromaDB
    print(f"\n  💾 Storing in ChromaDB...")
    t_store = time.time()
    
    _, collection = get_collection(collection_name)
    
    # Check for existing chunks with same tender_id (re-indexing)
    existing = collection.get(where={"tender_id": tender_id})
    if existing and existing["ids"]:
        print(f"    Removing {len(existing['ids'])} existing chunks for tender {tender_id}")
        collection.delete(ids=existing["ids"])
    
    # Insert in batches
    STORE_BATCH = 100
    for i in range(0, len(embedded_chunks), STORE_BATCH):
        batch = embedded_chunks[i:i + STORE_BATCH]
        
        ids = [c["id"] for c in batch]
        embeddings = [c["embedding"] for c in batch]
        documents = [c["text"] for c in batch]
        metadatas = [
            {
                "tender_id": c["tender_id"],
                "tender_name": c.get("tender_name", ""),
                "filename": c["filename"],
                "chunk_idx": c["chunk_idx"],
                "start_char": c["start_char"],
                "end_char": c["end_char"],
            }
            for c in batch
        ]
        
        collection.add(
            ids=ids,
            embeddings=embeddings,
            documents=documents,
            metadatas=metadatas,
        )
    
    stats["time_total_sec"] = round(time.time() - t_start, 2)
    
    # Summary
    print(f"\n{'─'*60}")
    print(f"  ✅ Индексация завершена:")
    print(f"     Файлов: {stats['files_processed']}")
    print(f"     Чанков: {stats['chunks_embedded']} / {stats['chunks_total']}")
    print(f"     Время: {stats['time_total_sec']} сек "
          f"(extract={stats['time_extract_sec']}s, embed={stats['time_embed_sec']}s)")
    print(f"     Стоимость: ~{stats['cost_rub']} ₽")
    print(f"     База: {CHROMA_DB_PATH}")
    print(f"{'─'*60}\n")
    
    return stats


def show_stats(collection_name: str = "tenders"):
    """Show ChromaDB collection statistics."""
    _, collection = get_collection(collection_name)
    
    count = collection.count()
    print(f"\n{'='*60}")
    print(f"ChromaDB Statistics — Collection: {collection_name}")
    print(f"{'='*60}")
    print(f"  Total chunks: {count}")
    
    if count > 0:
        # Get unique tender_ids
        all_meta = collection.get(include=["metadatas"])
        tender_ids = set()
        files = set()
        for m in all_meta["metadatas"]:
            tender_ids.add(m.get("tender_id", "unknown"))
            files.add(f"{m.get('tender_id', '?')}:{m.get('filename', '?')}")
        
        print(f"  Unique tenders: {len(tender_ids)}")
        print(f"  Unique files: {len(files)}")
        print(f"\n  Тендеры:")
        for tid in sorted(tender_ids):
            tid_chunks = sum(1 for m in all_meta["metadatas"] if m.get("tender_id") == tid)
            print(f"    • {tid}: {tid_chunks} chunks")
    
    print(f"\n  DB path: {CHROMA_DB_PATH}")
    print(f"{'='*60}\n")


# ═══════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description="RAG Indexer — индексация тендерных документов")
    parser.add_argument("--files", nargs="+", help="Файлы для индексации")
    parser.add_argument("--folder", help="Папка с файлами тендера")
    parser.add_argument("--tender-id", help="ID тендера (номер закупки)")
    parser.add_argument("--tender-name", default="", help="Название тендера")
    parser.add_argument("--stats", action="store_true", help="Показать статистику базы")
    parser.add_argument("--collection", default="tenders", help="Имя коллекции ChromaDB")
    
    args = parser.parse_args()
    
    if args.stats:
        show_stats(args.collection)
        return
    
    # Collect files
    files = []
    if args.files:
        files = args.files
    elif args.folder:
        folder = Path(args.folder)
        if folder.is_dir():
            for ext in ("*.pdf", "*.docx", "*.xlsx", "*.doc", "*.xls", "*.txt"):
                files.extend(str(f) for f in folder.glob(ext))
        else:
            print(f"[ERROR] Folder not found: {args.folder}")
            sys.exit(1)
    else:
        parser.print_help()
        sys.exit(1)
    
    if not files:
        print("[ERROR] No files found to index")
        sys.exit(1)
    
    tender_id = args.tender_id or "unknown"
    
    # Index
    stats = index_tender(
        files=files,
        tender_id=tender_id,
        tender_name=args.tender_name,
        collection_name=args.collection,
    )
    
    # Output stats as JSON
    print(json.dumps(stats, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
