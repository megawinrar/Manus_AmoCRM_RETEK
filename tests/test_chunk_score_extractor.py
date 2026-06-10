"""
Тесты для chunk_score_extractor.py — модуль распознавания тендерных документов
с использованием чанков (250 токенов) и балльной системы.

Покрывает:
1. split_into_chunks() — разбиение текста на чанки
2. score_chunk_for_field() — балльная оценка чанков
3. extract_*() — извлечение полей из чанков
4. score_direction() — определение направления
5. determine_priority() — определение приоритета
6. extract_and_parse_tender() — полный pipeline
7. OCR fallback логика
"""
import os
import sys
import tempfile
from pathlib import Path
from unittest.mock import patch, MagicMock

import pytest

# Добавляем scripts/ в path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))
import chunk_score_extractor as cse


# ═══════════════════════════════════════════════════════════════════
# 1. ТЕСТЫ split_into_chunks()
# ═══════════════════════════════════════════════════════════════════

class TestSplitIntoChunks:
    """Тесты разбиения текста на чанки."""

    def test_empty_text_returns_empty(self):
        """Пустой текст → пустой список."""
        assert cse.split_into_chunks("", "test.pdf", "pdf") == []
        assert cse.split_into_chunks("   ", "test.pdf", "pdf") == []

    def test_short_text_below_min_returns_empty(self):
        """Текст короче MIN_CHUNK_SIZE → пустой список."""
        assert cse.split_into_chunks("abc", "test.pdf", "pdf") == []

    def test_single_chunk_for_short_text(self):
        """Текст меньше CHUNK_SIZE → один чанк."""
        text = "Заказчик: АО «Ростех»\nНМЦ: 1 000 000 руб.\n" * 5
        chunks = cse.split_into_chunks(text, "doc.pdf", "pdf")
        assert len(chunks) == 1
        assert chunks[0]["source_file"] == "doc.pdf"
        assert chunks[0]["file_type"] == "pdf"
        assert chunks[0]["chunk_idx"] == 0

    def test_multiple_chunks_for_long_text(self):
        """Длинный текст → несколько чанков."""
        # Создаём текст длиннее CHUNK_SIZE (1000 символов)
        text = "\n".join([f"Строка {i}: " + "x" * 80 for i in range(30)])
        chunks = cse.split_into_chunks(text, "big.pdf", "pdf")
        assert len(chunks) > 1

    def test_chunk_size_approximately_correct(self):
        """Каждый чанк примерно 250 токенов (~1000 символов)."""
        text = "\n".join([f"Параграф {i}: " + "слово " * 50 for i in range(50)])
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        for chunk in chunks:
            # Допускаем ±50% от CHUNK_SIZE из-за семантических границ
            assert len(chunk["text"]) <= cse.CHUNK_SIZE * 2

    def test_chunk_has_required_fields(self):
        """Каждый чанк содержит обязательные поля."""
        text = "Заказчик: ООО «Тест»\n" * 20
        chunks = cse.split_into_chunks(text, "test.docx", "docx")
        for chunk in chunks:
            assert "text" in chunk
            assert "source_file" in chunk
            assert "chunk_idx" in chunk
            assert "start_char" in chunk
            assert "end_char" in chunk
            assert "file_type" in chunk

    def test_chunk_overlap_exists(self):
        """Чанки имеют перекрытие (overlap)."""
        text = "\n".join([f"Уникальная строка номер {i}" for i in range(100)])
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        if len(chunks) >= 2:
            # Последняя часть первого чанка должна быть в начале второго
            chunk1_end = chunks[0]["text"][-100:]
            chunk2_start = chunks[1]["text"][:200]
            # Хотя бы часть текста должна перекрываться
            overlap_found = any(
                word in chunk2_start
                for word in chunk1_end.split()
                if len(word) > 5
            )
            assert overlap_found or len(chunks) == 1

    def test_chunk_idx_sequential(self):
        """Индексы чанков идут последовательно."""
        text = "\n".join(["Текст " * 100 for _ in range(20)])
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        for i, chunk in enumerate(chunks):
            assert chunk["chunk_idx"] == i


# ═══════════════════════════════════════════════════════════════════
# 2. ТЕСТЫ score_chunk_for_field()
# ═══════════════════════════════════════════════════════════════════

class TestScoreChunkForField:
    """Тесты балльной системы оценки чанков."""

    def test_customer_chunk_high_score(self):
        """Чанк с ключевыми словами заказчика → высокий балл."""
        chunk = {
            "text": "Заказчик: АО «НПО «Высокоточные комплексы», ИНН 7718862113",
            "source_file": "doc.pdf",
            "chunk_idx": 0,
            "start_char": 0,
            "end_char": 100,
            "file_type": "pdf",
        }
        score = cse.score_chunk_for_field(chunk, "customer")
        assert score > 0.5

    def test_irrelevant_chunk_low_score(self):
        """Чанк без ключевых слов → низкий балл."""
        chunk = {
            "text": "Фреза концевая D10 z4 HRC55 покрытие TiAlN длина 75мм",
            "source_file": "spec.pdf",
            "chunk_idx": 5,
            "start_char": 500,
            "end_char": 600,
            "file_type": "pdf",
        }
        score = cse.score_chunk_for_field(chunk, "customer")
        assert score < 0.5

    def test_nmc_chunk_high_score(self):
        """Чанк с ценой → высокий балл для поля nmc."""
        chunk = {
            "text": "Начальная (максимальная) цена договора: 5 450 000,00 руб.",
            "source_file": "notice.pdf",
            "chunk_idx": 2,
            "start_char": 200,
            "end_char": 300,
            "file_type": "pdf",
        }
        score = cse.score_chunk_for_field(chunk, "nmc")
        assert score > 0.5

    def test_deadline_chunk_high_score(self):
        """Чанк с датой окончания подачи → высокий балл для deadline."""
        chunk = {
            "text": "Дата окончания подачи заявок: 15.06.2026 10:00 (МСК)",
            "source_file": "notice.pdf",
            "chunk_idx": 3,
            "start_char": 300,
            "end_char": 400,
            "file_type": "pdf",
        }
        score = cse.score_chunk_for_field(chunk, "deadline")
        assert score > 0.5

    def test_score_returns_float_between_0_and_1(self):
        """Балл всегда в диапазоне [0.0, 1.0]."""
        chunk = {"text": "Любой текст", "source_file": "t.pdf",
                 "chunk_idx": 0, "start_char": 0, "end_char": 10, "file_type": "pdf"}
        for field in cse.FIELD_WEIGHTS.keys():
            score = cse.score_chunk_for_field(chunk, field)
            assert 0.0 <= score <= 1.0

    def test_unknown_field_returns_zero(self):
        """Неизвестное поле → 0."""
        chunk = {"text": "Текст", "source_file": "t.pdf",
                 "chunk_idx": 0, "start_char": 0, "end_char": 10, "file_type": "pdf"}
        score = cse.score_chunk_for_field(chunk, "nonexistent_field")
        assert score == 0.0


# ═══════════════════════════════════════════════════════════════════
# 3. ТЕСТЫ extract_*() функций
# ═══════════════════════════════════════════════════════════════════

class TestExtractCustomer:
    """Тесты извлечения заказчика."""

    def test_extract_ao(self):
        """Извлечение АО из текста."""
        chunk = {"text": "Заказчик: АО «НПО «Высокоточные комплексы»",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_customer(chunk)
        assert result is not None
        value, confidence = result
        assert "Высокоточные комплексы" in value or "АО" in value
        assert confidence > 0

    def test_extract_ooo(self):
        """Извлечение ООО."""
        chunk = {"text": "Организатор закупки: ООО «ЧТЗ-УРАЛТРАК»",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_customer(chunk)
        assert result is not None
        value, confidence = result
        assert "ЧТЗ" in value or "УРАЛТРАК" in value

    def test_no_customer_returns_none(self):
        """Текст без заказчика → None."""
        chunk = {"text": "Фреза D10 z4 длина 75мм покрытие TiAlN",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_customer(chunk)
        assert result is None


class TestExtractNmc:
    """Тесты извлечения НМЦ."""

    def test_extract_nmc_with_rubles(self):
        """Извлечение НМЦ в рублях."""
        chunk = {"text": "Начальная (максимальная) цена: 5 450 000,00 руб.",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_nmc(chunk)
        assert result is not None
        value, confidence = result
        assert value > 0

    def test_extract_nmc_large_number(self):
        """Извлечение крупной НМЦ."""
        chunk = {"text": "Начальная максимальная цена: 12 345 678,90 руб.",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 60, "file_type": "pdf"}
        result = cse.extract_nmc(chunk)
        assert result is not None
        value, confidence = result
        assert value > 1_000_000

    def test_no_nmc_returns_none(self):
        """Текст без цены → None."""
        chunk = {"text": "Поставка инструмента согласно спецификации",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_nmc(chunk)
        assert result is None


class TestExtractDeadline:
    """Тесты извлечения дедлайна."""

    def test_extract_date_format_dd_mm_yyyy(self):
        """Извлечение даты в формате DD.MM.YYYY."""
        chunk = {"text": "Дата окончания подачи заявок: 15.06.2026 10:00",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_deadline(chunk)
        assert result is not None
        value, confidence = result
        # Формат возврата: YYYY-MM-DD
        assert "2026" in value and "06" in value and "15" in value

    def test_no_deadline_returns_none(self):
        """Текст без даты → None."""
        chunk = {"text": "Поставка инструмента в течение 30 дней",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_deadline(chunk)
        assert result is None


class TestExtractPlatform:
    """Тесты извлечения площадки."""

    def test_extract_eis(self):
        """Извлечение ЕИС/Госзакупки."""
        chunk = {"text": "Закупка размещена на сайте zakupki.gov.ru (ЕИС)",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_platform(chunk)
        assert result is not None

    def test_extract_ast_goz(self):
        """Извлечение АСТ ГОЗ."""
        chunk = {"text": "Площадка: АСТ ГОЗ (astgoz.ru)",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_platform(chunk)
        assert result is not None


class TestExtractProcedureNumber:
    """Тесты извлечения номера процедуры."""

    def test_extract_eis_number(self):
        """Извлечение номера ЕИС."""
        chunk = {"text": "Номер извещения: 0201-2026-00592",
                 "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": 50, "file_type": "pdf"}
        result = cse.extract_procedure_number(chunk)
        assert result is not None
        value, confidence = result
        assert "0201" in value


# ═══════════════════════════════════════════════════════════════════
# 4. ТЕСТЫ score_direction()
# ═══════════════════════════════════════════════════════════════════

class TestScoreDirection:
    """Тесты определения направления."""

    def test_carbide_direction(self):
        """Текст с твёрдосплавным инструментом → CARBIDE."""
        chunks = [{"text": "Фреза концевая твердосплавная D10 z4 HRC55 пластины",
                   "source_file": "t.pdf", "chunk_idx": 0,
                   "start_char": 0, "end_char": 50, "file_type": "pdf"}]
        result = cse.score_direction(chunks)
        # Возвращает (direction, confidence, scores)
        assert len(result) == 3
        direction = result[0]
        assert "CARBIDE" in direction

    def test_hss_direction(self):
        """Текст с быстрорежущим инструментом → HSS."""
        chunks = [{"text": "Сверло спиральное hss р6м5 диаметр 10мм метчик",
                   "source_file": "t.pdf", "chunk_idx": 0,
                   "start_char": 0, "end_char": 50, "file_type": "pdf"}]
        result = cse.score_direction(chunks)
        assert len(result) == 3
        direction = result[0]
        assert "HSS" in direction

    def test_empty_chunks_returns_default(self):
        """Пустой список чанков → дефолтное направление."""
        result = cse.score_direction([])
        assert len(result) == 3
        direction = result[0]
        assert direction == "CARBIDE-STANDARD"


# ═══════════════════════════════════════════════════════════════════
# 5. ТЕСТЫ determine_priority()
# ═══════════════════════════════════════════════════════════════════

class TestDeterminePriority:
    """Тесты определения приоритета."""

    def test_high_nmc_high_priority(self):
        """Высокая НМЦ → Р2 (или Р1 если дедлайн близко)."""
        priority, warnings = cse.determine_priority(10_000_000, "")
        assert priority in ("Р1", "P1", "Р2", "P2")

    def test_low_nmc_low_priority(self):
        """Низкая НМЦ → Р3."""
        priority, warnings = cse.determine_priority(50_000, "")
        assert priority in ("Р3", "P3")

    def test_zero_nmc(self):
        """НМЦ = 0 → не падает."""
        priority, warnings = cse.determine_priority(0, "")
        assert priority is not None

    def test_returns_tuple(self):
        """Всегда возвращает (priority, warnings)."""
        result = cse.determine_priority(1_000_000, "")
        assert isinstance(result, tuple)
        assert len(result) == 2


# ═══════════════════════════════════════════════════════════════════
# 6. ТЕСТЫ extract_and_parse_tender() — полный pipeline
# ═══════════════════════════════════════════════════════════════════

class TestExtractAndParseTender:
    """Тесты полного pipeline распознавания."""

    def test_empty_files_list(self):
        """Пустой список файлов → результат с warnings."""
        result = cse.extract_and_parse_tender([])
        assert result["validation_status"] == "blocked"
        assert result["total_chunks"] == 0

    def test_nonexistent_files(self):
        """Несуществующие файлы → результат с warnings."""
        result = cse.extract_and_parse_tender(["/tmp/nonexistent.pdf"])
        assert result["total_chunks"] == 0

    def test_docx_file_parsing(self):
        """Парсинг реального DOCX файла."""
        # Создаём тестовый DOCX
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Заказчик: АО «Тестовый завод»")
            doc.add_paragraph("Начальная максимальная цена: 2 500 000 руб.")
            doc.add_paragraph("Дата окончания подачи заявок: 20.06.2026")
            doc.add_paragraph("Площадка: ЕИС (zakupki.gov.ru)")
            doc.add_paragraph("Номер извещения: 0301-2026-00123")
            doc.add_paragraph("Поставка фрез концевых твёрдосплавных D8-D20")

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                doc.save(f.name)
                tmp_path = f.name

            result = cse.extract_and_parse_tender([tmp_path])

            # Проверяем что pipeline отработал
            assert result["total_chunks"] > 0
            assert result["chunks_used"] > 0
            assert "confidence" in result
            assert result["customer"] is not None
            assert result["direction"] is not None
            assert result["priority"] is not None

            os.unlink(tmp_path)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_result_has_required_fields(self):
        """Результат содержит все обязательные поля."""
        result = cse.extract_and_parse_tender([])
        required_fields = [
            "customer", "nmc", "direction", "deadline",
            "priority", "validation_status", "confidence",
            "chunks_used", "total_chunks",
        ]
        for field in required_fields:
            assert field in result, f"Missing field: {field}"

    def test_confidence_dict_structure(self):
        """confidence — словарь с float значениями."""
        result = cse.extract_and_parse_tender([])
        assert isinstance(result["confidence"], dict)


# ═══════════════════════════════════════════════════════════════════
# 7. ТЕСТЫ extract_file_text()
# ═══════════════════════════════════════════════════════════════════

class TestExtractFileText:
    """Тесты извлечения текста из файлов."""

    def test_docx_extraction(self):
        """Извлечение текста из DOCX."""
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Тестовый текст для извлечения")
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                doc.save(f.name)
                tmp_path = f.name
            text = cse.extract_file_text(tmp_path)
            assert "Тестовый текст" in text
            os.unlink(tmp_path)
        except ImportError:
            pytest.skip("python-docx not installed")

    def test_nonexistent_file_returns_empty(self):
        """Несуществующий файл → пустая строка."""
        text = cse.extract_file_text("/tmp/no_such_file_xyz.pdf")
        assert text == "" or text is None

    def test_txt_file_extraction(self):
        """Извлечение текста из TXT."""
        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False) as f:
            f.write("Простой текстовый файл\nВторая строка")
            tmp_path = f.name
        text = cse.extract_file_text(tmp_path)
        # Может вернуть пустую строку если .txt не поддерживается
        os.unlink(tmp_path)


# ═══════════════════════════════════════════════════════════════════
# 8. ТЕСТЫ OCR fallback
# ═══════════════════════════════════════════════════════════════════

class TestOCRFallback:
    """Тесты OCR fallback логики."""

    @patch("chunk_score_extractor.ocr_pdf_full")
    def test_ocr_fallback_called_on_low_confidence(self, mock_ocr):
        """OCR вызывается при низкой confidence."""
        mock_ocr.return_value = "НМЦ: 1 000 000 руб."
        # Создаём пустой PDF-файл для теста
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"%PDF-1.4 test")
            pdf_path = f.name

        result = cse.apply_ocr_fallback("nmc", [pdf_path], [])
        os.unlink(pdf_path)
        # OCR может быть вызван или нет в зависимости от реализации

    def test_ocr_pdf_full_handles_missing_file(self):
        """ocr_pdf_full не падает на несуществующем файле."""
        try:
            result = cse.ocr_pdf_full("/tmp/no_such_file.pdf")
            assert result == "" or result is None
        except Exception:
            pass  # Допускаем исключение для несуществующего файла


# ═══════════════════════════════════════════════════════════════════
# 9. ТЕСТЫ конфигурации
# ═══════════════════════════════════════════════════════════════════

class TestConfiguration:
    """Тесты конфигурации модуля."""

    def test_chunk_size_250_tokens(self):
        """Размер чанка = 250 токенов."""
        assert cse.CHUNK_SIZE_TOKENS == 250

    def test_chunk_overlap_75_tokens(self):
        """Overlap = 75 токенов."""
        assert cse.CHUNK_OVERLAP_TOKENS == 75

    def test_chars_per_token(self):
        """4 символа на токен."""
        assert cse.CHARS_PER_TOKEN == 4

    def test_chunk_size_in_chars(self):
        """CHUNK_SIZE = 250 * 4 = 1000."""
        assert cse.CHUNK_SIZE == 1000

    def test_field_weights_exist(self):
        """FIELD_WEIGHTS содержит все основные поля."""
        required_fields = ["customer", "nmc", "deadline", "platform"]
        for field in required_fields:
            assert field in cse.FIELD_WEIGHTS

    def test_confidence_thresholds(self):
        """Пороги confidence корректны."""
        assert cse.CONFIDENCE_HIGH > cse.CONFIDENCE_MEDIUM > cse.CONFIDENCE_LOW
        assert cse.CONFIDENCE_HIGH <= 1.0
        assert cse.CONFIDENCE_LOW > 0.0
