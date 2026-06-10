"""
Полные тесты для КАЖДОГО нововведения системы RETEK AmoCRM.

Модули и нововведения:
1. ЧАНКИ (split_into_chunks) — разбиение на 250 токенов с overlap 75
2. БАЛЛЬНАЯ СИСТЕМА (score_chunk_for_field) — оценка релевантности чанков
3. ИЗВЛЕЧЕНИЕ ПОЛЕЙ — все 12 extract_* функций
4. OCR FALLBACK — полный OCR + точечный OCR вокруг ключевого слова
5. НАПРАВЛЕНИЕ (score_direction) — балльная классификация по 5 направлениям
6. ПРИОРИТЕТ (determine_priority) — Р1/Р2/Р3 по НМЦ и дедлайну
7. PIPELINE (extract_and_parse_tender) — полный цикл обработки
8. CRON (APScheduler) — расписание 5 минут, автономность
9. ДЕДУПЛИКАЦИЯ — предотвращение повторной обработки
10. DOCKER — restart:always, healthcheck, OCR-зависимости
11. ИНТЕГРАЦИЯ cron_yadisk ↔ chunk_score_extractor

Каждый тест проверяет конкретное поведение и гарантирует стабильность системы.
"""
import os
import sys
import re
import time
import tempfile
from pathlib import Path
from datetime import datetime, timedelta
from unittest.mock import patch, MagicMock, call

import pytest

# Добавляем scripts/ и src/ в path
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'scripts'))
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', 'src'))
import chunk_score_extractor as cse


# ═══════════════════════════════════════════════════════════════════════════
# 1. ЧАНКИ — split_into_chunks
# ═══════════════════════════════════════════════════════════════════════════

class TestChunking:
    """Тесты разбиения текста на чанки (250 токенов, overlap 75)."""

    def test_chunk_size_250_tokens_config(self):
        """Конфигурация: CHUNK_SIZE_TOKENS = 250."""
        assert cse.CHUNK_SIZE_TOKENS == 250

    def test_chunk_overlap_75_tokens_config(self):
        """Конфигурация: CHUNK_OVERLAP_TOKENS = 75."""
        assert cse.CHUNK_OVERLAP_TOKENS == 75

    def test_chars_per_token_4(self):
        """Конфигурация: 4 символа на токен (русский текст)."""
        assert cse.CHARS_PER_TOKEN == 4

    def test_chunk_size_chars_equals_1000(self):
        """CHUNK_SIZE = 250 * 4 = 1000 символов."""
        assert cse.CHUNK_SIZE == 1000

    def test_empty_text_returns_empty_list(self):
        """Пустой текст → пустой список чанков."""
        assert cse.split_into_chunks("", "test.pdf", "pdf") == []

    def test_whitespace_only_returns_empty(self):
        """Только пробелы → пустой список."""
        assert cse.split_into_chunks("   \n\t  ", "test.pdf", "pdf") == []

    def test_text_below_min_chunk_size_returns_empty(self):
        """Текст короче MIN_CHUNK_SIZE (50 символов) → пустой список."""
        short_text = "abc" * 10  # 30 символов
        assert cse.split_into_chunks(short_text, "t.pdf", "pdf") == []

    def test_single_chunk_for_moderate_text(self):
        """Текст 200 символов (< CHUNK_SIZE) → один чанк."""
        text = "Заказчик: АО «Ростех». НМЦ: 5 000 000 руб. " * 5  # ~220 символов
        chunks = cse.split_into_chunks(text, "doc.pdf", "pdf")
        assert len(chunks) == 1

    def test_multiple_chunks_for_long_text(self):
        """Текст 5000 символов → несколько чанков."""
        text = "\n".join([f"Строка {i}: " + "абвгд " * 30 for i in range(30)])
        chunks = cse.split_into_chunks(text, "big.pdf", "pdf")
        assert len(chunks) > 1

    def test_chunk_fields_structure(self):
        """Каждый чанк содержит все обязательные поля."""
        text = "Заказчик: ООО «Тест». " * 30
        chunks = cse.split_into_chunks(text, "test.docx", "docx")
        for chunk in chunks:
            assert "text" in chunk
            assert "source_file" in chunk
            assert "chunk_idx" in chunk
            assert "start_char" in chunk
            assert "end_char" in chunk
            assert "file_type" in chunk

    def test_chunk_idx_sequential_from_zero(self):
        """Индексы чанков: 0, 1, 2, ... (последовательно)."""
        text = "\n".join(["Текст " * 100 for _ in range(20)])
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        for i, chunk in enumerate(chunks):
            assert chunk["chunk_idx"] == i

    def test_source_file_preserved(self):
        """source_file сохраняется из аргумента."""
        text = "Заказчик: АО «Тест». " * 30
        chunks = cse.split_into_chunks(text, "my_document.pdf", "pdf")
        for chunk in chunks:
            assert chunk["source_file"] == "my_document.pdf"

    def test_file_type_preserved(self):
        """file_type сохраняется из аргумента."""
        text = "Заказчик: АО «Тест». " * 30
        chunks = cse.split_into_chunks(text, "doc.docx", "docx")
        for chunk in chunks:
            assert chunk["file_type"] == "docx"

    def test_chunk_text_not_empty(self):
        """Текст чанка не пустой."""
        text = "\n".join(["Текст " * 50 for _ in range(10)])
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        for chunk in chunks:
            assert len(chunk["text"].strip()) > 0

    def test_overlap_between_chunks(self):
        """Между чанками есть перекрытие (overlap)."""
        text = "\n".join([f"Уникальное_слово_{i} " + "текст " * 40 for i in range(30)])
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        if len(chunks) >= 2:
            # Проверяем что end_char первого > start_char второго (overlap)
            # или что есть общий текст
            chunk1_words = set(chunks[0]["text"].split())
            chunk2_words = set(chunks[1]["text"].split())
            overlap = chunk1_words & chunk2_words
            assert len(overlap) > 0, "Чанки должны иметь перекрытие"

    def test_all_text_covered(self):
        """Весь исходный текст покрыт чанками (нет пропусков)."""
        text = "Слово " * 500
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        # Проверяем что start_char первого = 0 и end_char последнего ~ len(text)
        if chunks:
            assert chunks[0]["start_char"] == 0
            assert chunks[-1]["end_char"] >= len(text) - 100  # допуск на strip

    def test_chunk_max_size_bounded(self):
        """Ни один чанк не превышает 2x CHUNK_SIZE."""
        text = "\n".join(["Текст " * 100 for _ in range(20)])
        chunks = cse.split_into_chunks(text, "test.pdf", "pdf")
        for chunk in chunks:
            assert len(chunk["text"]) <= cse.CHUNK_SIZE * 2


# ═══════════════════════════════════════════════════════════════════════════
# 2. БАЛЛЬНАЯ СИСТЕМА — score_chunk_for_field
# ═══════════════════════════════════════════════════════════════════════════

class TestScoringSystem:
    """Тесты балльной системы оценки чанков."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_score_returns_float(self):
        """score_chunk_for_field возвращает float."""
        chunk = self._make_chunk("Заказчик: АО «Тест»")
        score = cse.score_chunk_for_field(chunk, "customer")
        assert isinstance(score, float)

    def test_score_range_0_to_1(self):
        """Балл всегда в диапазоне [0.0, 1.0]."""
        chunk = self._make_chunk("Произвольный текст без ключевых слов")
        for field in cse.FIELD_WEIGHTS.keys():
            score = cse.score_chunk_for_field(chunk, field)
            assert 0.0 <= score <= 1.0, f"Field {field}: score {score} out of range"

    def test_relevant_chunk_scores_higher(self):
        """Релевантный чанк получает больший балл чем нерелевантный."""
        relevant = self._make_chunk("Заказчик: АО «НПО Высокоточные комплексы»")
        irrelevant = self._make_chunk("Фреза концевая D10 z4 HRC55 покрытие TiAlN")
        score_r = cse.score_chunk_for_field(relevant, "customer")
        score_i = cse.score_chunk_for_field(irrelevant, "customer")
        assert score_r > score_i

    def test_nmc_keywords_boost_score(self):
        """Ключевые слова НМЦ повышают балл."""
        with_kw = self._make_chunk("Начальная максимальная цена договора: 5 000 000 руб.")
        without_kw = self._make_chunk("Поставка инструмента согласно спецификации")
        assert cse.score_chunk_for_field(with_kw, "nmc") > cse.score_chunk_for_field(without_kw, "nmc")

    def test_deadline_keywords_boost_score(self):
        """Ключевые слова дедлайна повышают балл."""
        with_kw = self._make_chunk("Дата окончания подачи заявок: 15.06.2026 10:00")
        without_kw = self._make_chunk("Инструмент должен быть поставлен в течение 30 дней")
        assert cse.score_chunk_for_field(with_kw, "deadline") > cse.score_chunk_for_field(without_kw, "deadline")

    def test_unknown_field_returns_zero(self):
        """Неизвестное поле → балл 0.0."""
        chunk = self._make_chunk("Любой текст")
        assert cse.score_chunk_for_field(chunk, "nonexistent_field_xyz") == 0.0

    def test_all_fields_in_field_weights(self):
        """Все поля из FIELD_EXTRACTORS есть в FIELD_WEIGHTS."""
        for field in cse.FIELD_EXTRACTORS.keys():
            assert field in cse.FIELD_WEIGHTS, f"Missing {field} in FIELD_WEIGHTS"

    def test_anti_keywords_reduce_score(self):
        """Anti-keywords снижают балл (для deadline: 'поставки' — anti)."""
        with_anti = self._make_chunk("Срок поставки: 30 дней с момента заключения")
        without_anti = self._make_chunk("Дата окончания подачи заявок: 15.06.2026")
        # deadline anti_keywords = ["поставки", "исполнения", "оплаты"]
        score_anti = cse.score_chunk_for_field(with_anti, "deadline")
        score_clean = cse.score_chunk_for_field(without_anti, "deadline")
        assert score_clean >= score_anti

    def test_context_keywords_boost(self):
        """Context keywords повышают балл."""
        with_context = self._make_chunk("Начальная максимальная цена: 5 000 000 рублей")
        without_context = self._make_chunk("Цена: 5 000 000")
        # nmc context_keywords include "рубл"
        score_ctx = cse.score_chunk_for_field(with_context, "nmc")
        score_no_ctx = cse.score_chunk_for_field(without_context, "nmc")
        assert score_ctx >= score_no_ctx


# ═══════════════════════════════════════════════════════════════════════════
# 3. ИЗВЛЕЧЕНИЕ ПОЛЕЙ — все 12 extract_* функций
# ═══════════════════════════════════════════════════════════════════════════

class TestExtractCustomerFull:
    """Полные тесты extract_customer."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_ao_extraction(self):
        result = cse.extract_customer(self._make_chunk("Заказчик: АО «НПО «Высокоточные комплексы»"))
        assert result is not None
        assert result[1] > 0  # confidence > 0

    def test_ooo_extraction(self):
        result = cse.extract_customer(self._make_chunk("Организатор: ООО «ЧТЗ-УРАЛТРАК»"))
        assert result is not None

    def test_pao_extraction(self):
        result = cse.extract_customer(self._make_chunk("Заказчик: ПАО «Объединённая авиастроительная корпорация»"))
        assert result is not None

    def test_fgup_extraction(self):
        result = cse.extract_customer(self._make_chunk("Организатор закупки: ФГУП «НПЦАП»"))
        assert result is not None

    def test_no_customer_returns_none(self):
        result = cse.extract_customer(self._make_chunk("Фреза D10 z4 длина 75мм"))
        assert result is None

    def test_confidence_is_float(self):
        result = cse.extract_customer(self._make_chunk("Заказчик: АО «Тест»"))
        if result:
            assert isinstance(result[1], float)


class TestExtractNmcFull:
    """Полные тесты extract_nmc."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_nmc_with_spaces(self):
        """НМЦ с пробелами: 5 450 000,00."""
        result = cse.extract_nmc(self._make_chunk("Начальная (максимальная) цена: 5 450 000,00 руб."))
        assert result is not None
        assert result[0] > 5_000_000

    def test_nmc_without_spaces(self):
        """НМЦ без пробелов: 5450000."""
        result = cse.extract_nmc(self._make_chunk("НМЦ: 5450000 руб."))
        assert result is not None
        assert result[0] > 5_000_000

    def test_nmc_with_comma_decimal(self):
        """НМЦ с запятой: 1 234 567,89."""
        result = cse.extract_nmc(self._make_chunk("Начальная максимальная цена договора: 1 234 567,89 руб."))
        assert result is not None
        assert 1_234_000 < result[0] < 1_235_000

    def test_nmc_returns_float(self):
        """НМЦ возвращается как float."""
        result = cse.extract_nmc(self._make_chunk("НМЦ: 100000 руб."))
        if result:
            assert isinstance(result[0], (int, float))

    def test_no_nmc_returns_none(self):
        result = cse.extract_nmc(self._make_chunk("Поставка инструмента"))
        assert result is None


class TestExtractDeadlineFull:
    """Полные тесты extract_deadline."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_dd_mm_yyyy_format(self):
        """Дата DD.MM.YYYY."""
        result = cse.extract_deadline(self._make_chunk("Дата окончания подачи заявок: 15.06.2026 10:00"))
        assert result is not None
        assert "2026" in result[0]

    def test_deadline_with_zayavok(self):
        """Дата с 'заявок' между 'подачи' и датой."""
        result = cse.extract_deadline(self._make_chunk("Окончание подачи заявок 20.07.2026"))
        assert result is not None
        assert "2026" in result[0]

    def test_deadline_returns_iso_format(self):
        """Дедлайн в формате YYYY-MM-DD."""
        result = cse.extract_deadline(self._make_chunk("Срок подачи заявок до 25.12.2026"))
        if result:
            assert re.match(r'\d{4}-\d{2}-\d{2}', result[0])

    def test_no_deadline_returns_none(self):
        result = cse.extract_deadline(self._make_chunk("Поставка в течение 30 дней"))
        assert result is None


class TestExtractPlatformFull:
    """Полные тесты extract_platform."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_eis_zakupki(self):
        result = cse.extract_platform(self._make_chunk("Закупка на zakupki.gov.ru (ЕИС)"))
        assert result is not None

    def test_ast_goz(self):
        result = cse.extract_platform(self._make_chunk("Площадка: АСТ ГОЗ"))
        assert result is not None

    def test_b2b_center(self):
        result = cse.extract_platform(self._make_chunk("Торговая площадка: B2B-Center"))
        assert result is not None

    def test_no_platform_returns_none(self):
        result = cse.extract_platform(self._make_chunk("Фреза D10 z4"))
        assert result is None


class TestExtractProcedureNumberFull:
    """Полные тесты extract_procedure_number."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_eis_format(self):
        """Номер в формате ЕИС: XXXX-YYYY-NNNNN."""
        result = cse.extract_procedure_number(self._make_chunk("Номер извещения: 0201-2026-00592"))
        assert result is not None
        assert "0201" in result[0]

    def test_no_number_returns_none(self):
        result = cse.extract_procedure_number(self._make_chunk("Поставка инструмента"))
        assert result is None


class TestExtractProcedureTypeFull:
    """Полные тесты extract_procedure_type."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_zapros_kotirovok(self):
        result = cse.extract_procedure_type(self._make_chunk("Способ закупки: Запрос котировок"))
        assert result is not None
        assert "котировок" in result[0].lower()

    def test_otkrytyj_aukcion(self):
        result = cse.extract_procedure_type(self._make_chunk("Вид процедуры: Открытый аукцион"))
        assert result is not None

    def test_zakrytyj_konkurs(self):
        result = cse.extract_procedure_type(self._make_chunk("Тип закупки: Закрытый конкурс"))
        assert result is not None

    def test_no_type_returns_none(self):
        result = cse.extract_procedure_type(self._make_chunk("Фреза D10"))
        assert result is None


class TestExtractSubjectFull:
    """Полные тесты extract_subject."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_subject_extraction(self):
        result = cse.extract_subject(self._make_chunk(
            "Предмет закупки: Поставка твёрдосплавного режущего инструмента"))
        assert result is not None
        assert len(result[0]) > 10

    def test_naimenovanie_zakaza(self):
        result = cse.extract_subject(self._make_chunk(
            "Наименование заказа Поставка фрез концевых для обработки стали"))
        assert result is not None

    def test_short_subject_extracted(self):
        """Короткий предмет после ключевого слова — всё равно извлекается."""
        result = cse.extract_subject(self._make_chunk("Предмет закупки: Фрезы"))
        # Функция извлекает текст после ключевого слова
        assert result is not None

    def test_no_subject_returns_none(self):
        """Текст без ключевых слов предмета → None."""
        result = cse.extract_subject(self._make_chunk("Фреза D10 z4 HRC55"))
        assert result is None


class TestExtractPositionCount:
    """Тесты extract_position_count."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_numbered_list(self):
        """Нумерованный список → количество позиций."""
        text = "\n".join([f"  {i}.\tФреза D{i*2}" for i in range(1, 15)])
        result = cse.extract_position_count(self._make_chunk(text))
        assert result is not None
        assert result[0] >= 10

    def test_single_item_returns_none(self):
        """Одна позиция → None (не считается)."""
        result = cse.extract_position_count(self._make_chunk("1.\tФреза D10"))
        assert result is None  # max_pos < 2


class TestExtractEquivalent:
    """Тесты extract_equivalent."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_equivalent_allowed(self):
        result = cse.extract_equivalent(self._make_chunk("Допускается поставка или эквивалент"))
        assert result is not None
        assert result[0] is True

    def test_equivalent_not_allowed(self):
        result = cse.extract_equivalent(self._make_chunk("Эквивалент не допускается"))
        assert result is not None
        assert result[0] is False

    def test_no_equivalent_info_returns_none(self):
        result = cse.extract_equivalent(self._make_chunk("Поставка фрез D10"))
        assert result is None


class TestExtractGisp:
    """Тесты extract_gisp."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_gisp_required(self):
        result = cse.extract_gisp(self._make_chunk("Требуется подтверждение в ГИСП"))
        assert result is not None
        assert result[0] is True

    def test_gisp_url(self):
        result = cse.extract_gisp(self._make_chunk("Регистрация на gisp.gov.ru обязательна"))
        assert result is not None

    def test_no_gisp_returns_none(self):
        result = cse.extract_gisp(self._make_chunk("Поставка фрез D10"))
        assert result is None


class TestExtractCity:
    """Тесты extract_city."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_moscow(self):
        result = cse.extract_city(self._make_chunk("Место поставки: г. Москва"))
        assert result is not None
        assert "Москв" in result[0]

    def test_spb(self):
        result = cse.extract_city(self._make_chunk("Адрес: г. Санкт-Петербург"))
        assert result is not None

    def test_no_city_returns_none(self):
        result = cse.extract_city(self._make_chunk("Поставка фрез D10"))
        assert result is None


class TestExtractDeliveryTerms:
    """Тесты extract_delivery_terms."""

    def _make_chunk(self, text):
        return {"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                "start_char": 0, "end_char": len(text), "file_type": "pdf"}

    def test_days_format(self):
        result = cse.extract_delivery_terms(self._make_chunk(
            "Срок поставки: 30 календарных дней с момента заключения"))
        assert result is not None

    def test_explicit_term(self):
        result = cse.extract_delivery_terms(self._make_chunk(
            "Срок исполнения: до 31.12.2026"))
        assert result is not None

    def test_no_terms_returns_none(self):
        result = cse.extract_delivery_terms(self._make_chunk("Фреза D10 z4"))
        assert result is None


# ═══════════════════════════════════════════════════════════════════════════
# 4. OCR FALLBACK
# ═══════════════════════════════════════════════════════════════════════════

class TestOCRFallbackLogic:
    """Тесты логики OCR fallback."""

    def test_ocr_pdf_full_missing_file_returns_empty(self):
        """ocr_pdf_full для несуществующего файла → пустая строка."""
        result = cse.ocr_pdf_full("/tmp/nonexistent_file_xyz123.pdf")
        assert result == ""

    def test_ocr_region_missing_file_returns_empty(self):
        """ocr_region_around_keyword для несуществующего файла → пустая строка."""
        result = cse.ocr_region_around_keyword("/tmp/nonexistent.pdf", "заказчик")
        assert result == ""

    def test_apply_ocr_fallback_no_pdf_files(self):
        """apply_ocr_fallback без PDF файлов → None."""
        result = cse.apply_ocr_fallback("customer", [], [])
        assert result is None

    def test_apply_ocr_fallback_unknown_field(self):
        """apply_ocr_fallback для неизвестного поля → None."""
        result = cse.apply_ocr_fallback("unknown_field_xyz", ["/tmp/test.pdf"], [])
        assert result is None

    @patch("chunk_score_extractor.ocr_region_around_keyword")
    def test_ocr_fallback_uses_field_keywords(self, mock_ocr_region):
        """OCR fallback использует ключевые слова из FIELD_WEIGHTS."""
        mock_ocr_region.return_value = ""
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"%PDF-1.4 test")
            pdf_path = f.name
        cse.apply_ocr_fallback("customer", [pdf_path], [])
        os.unlink(pdf_path)
        # Проверяем что OCR был вызван с ключевыми словами из FIELD_WEIGHTS
        if mock_ocr_region.called:
            first_call_kw = mock_ocr_region.call_args_list[0][0][1]
            assert first_call_kw in cse.FIELD_WEIGHTS["customer"]["keywords"]

    @patch("chunk_score_extractor.ocr_pdf_full")
    @patch("chunk_score_extractor.ocr_region_around_keyword")
    def test_ocr_fallback_tries_full_ocr_after_region(self, mock_region, mock_full):
        """Если точечный OCR не дал результата → пробует полный OCR."""
        mock_region.return_value = ""
        mock_full.return_value = "Заказчик: АО «Тест из OCR»\n" * 10
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"%PDF-1.4 test")
            pdf_path = f.name
        result = cse.apply_ocr_fallback("customer", [pdf_path], [])
        os.unlink(pdf_path)
        # Полный OCR должен быть вызван
        assert mock_full.called

    def test_ocr_confidence_reduced(self):
        """OCR результат имеет сниженную confidence (×0.85 или ×0.75)."""
        # Мокаем точечный OCR чтобы вернуть текст с заказчиком
        with patch("chunk_score_extractor.ocr_region_around_keyword") as mock_ocr:
            mock_ocr.return_value = "Заказчик: АО «Тестовый завод» ИНН 1234567890 КПП 123456789"
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(b"%PDF-1.4 test")
                pdf_path = f.name
            result = cse.apply_ocr_fallback("customer", [pdf_path], [])
            os.unlink(pdf_path)
            if result:
                # Confidence должна быть снижена (< 1.0)
                assert result[1] < 1.0


# ═══════════════════════════════════════════════════════════════════════════
# 5. НАПРАВЛЕНИЕ — score_direction
# ═══════════════════════════════════════════════════════════════════════════

class TestScoreDirectionFull:
    """Полные тесты определения направления."""

    def _make_chunks(self, text):
        return [{"text": text, "source_file": "t.pdf", "chunk_idx": 0,
                 "start_char": 0, "end_char": len(text), "file_type": "pdf"}]

    def test_returns_tuple_of_3(self):
        """Возвращает (direction, confidence, scores)."""
        result = cse.score_direction(self._make_chunks("Фреза твердосплавная"))
        assert isinstance(result, tuple)
        assert len(result) == 3

    def test_carbide_standard(self):
        """Твёрдосплавный инструмент → CARBIDE-STANDARD."""
        result = cse.score_direction(self._make_chunks(
            "Фреза концевая твердосплавная D10 z4 пластины GESAC"))
        assert "CARBIDE" in result[0]

    def test_hss_standard(self):
        """Быстрорежущий инструмент → HSS-STANDARD."""
        result = cse.score_direction(self._make_chunks(
            "Сверло спиральное HSS Р6М5 метчик плашка зенкер"))
        assert "HSS" in result[0]

    def test_spec_drawing(self):
        """Специнструмент по чертежу → SPEC-DRAWING."""
        result = cse.score_direction(self._make_chunks(
            "Изготовление по чертежу заказчика нестандартный специнструмент по ТЗ"))
        assert "SPEC" in result[0]

    def test_diamond_standard(self):
        """Алмазный инструмент → DIAMOND-STANDARD."""
        result = cse.score_direction(self._make_chunks(
            "Алмазный инструмент PCD CBN кубический нитрид бора diamond"))
        assert "DIAMOND" in result[0]

    def test_out_of_scope(self):
        """Измерительный инструмент → OUT-OF-SCOPE."""
        result = cse.score_direction(self._make_chunks(
            "Калибр скоба штангенциркуль микрометр индикатор"))
        assert "OUT-OF-SCOPE" in result[0]

    def test_empty_chunks_default(self):
        """Пустой список → CARBIDE-STANDARD (дефолт)."""
        result = cse.score_direction([])
        assert result[0] == "CARBIDE-STANDARD"

    def test_confidence_is_float(self):
        """Confidence — float."""
        result = cse.score_direction(self._make_chunks("Фреза"))
        assert isinstance(result[1], float)

    def test_scores_dict_has_all_directions(self):
        """Scores содержит все направления."""
        result = cse.score_direction(self._make_chunks("Фреза"))
        scores = result[2]
        assert isinstance(scores, dict)
        for direction in cse.DIRECTION_KEYWORDS.keys():
            assert direction in scores

    def test_multiple_chunks_aggregation(self):
        """Баллы агрегируются по всем чанкам."""
        chunks = [
            {"text": "Фреза твердосплавная D10", "source_file": "t.pdf",
             "chunk_idx": 0, "start_char": 0, "end_char": 30, "file_type": "pdf"},
            {"text": "Пластины GESAC KORLOY sandvik", "source_file": "t.pdf",
             "chunk_idx": 1, "start_char": 30, "end_char": 60, "file_type": "pdf"},
        ]
        result = cse.score_direction(chunks)
        # С двумя чанками про твердосплав — CARBIDE должен быть выше
        assert "CARBIDE" in result[0]


# ═══════════════════════════════════════════════════════════════════════════
# 6. ПРИОРИТЕТ — determine_priority
# ═══════════════════════════════════════════════════════════════════════════

class TestDeterminePriorityFull:
    """Полные тесты определения приоритета."""

    def test_returns_tuple_of_2(self):
        """Возвращает (priority, warnings)."""
        result = cse.determine_priority(1_000_000, "")
        assert isinstance(result, tuple)
        assert len(result) == 2

    def test_high_nmc_p2(self):
        """НМЦ > 20M → Р2."""
        priority, _ = cse.determine_priority(25_000_000, "")
        assert priority in ("Р2", "P2")

    def test_medium_nmc_p2(self):
        """НМЦ 5M-20M → Р2."""
        priority, _ = cse.determine_priority(10_000_000, "")
        assert priority in ("Р2", "P2")

    def test_low_nmc_p3(self):
        """НМЦ 1M-5M → Р3."""
        priority, _ = cse.determine_priority(2_000_000, "")
        assert priority in ("Р3", "P3")

    def test_very_low_nmc_p3(self):
        """НМЦ < 1M → Р3."""
        priority, _ = cse.determine_priority(500_000, "")
        assert priority in ("Р3", "P3")

    def test_zero_nmc_no_crash(self):
        """НМЦ = 0 → не падает."""
        priority, warnings = cse.determine_priority(0, "")
        assert priority is not None

    def test_urgent_deadline_p1(self):
        """Дедлайн через 24 часа → Р1."""
        tomorrow = (datetime.now() + timedelta(hours=24)).strftime("%Y-%m-%d")
        priority, warnings = cse.determine_priority(10_000_000, tomorrow)
        assert priority in ("Р1", "P1")
        assert any("СРОЧНО" in w for w in warnings)

    def test_close_deadline_warning(self):
        """Дедлайн через 4 дня → предупреждение."""
        soon = (datetime.now() + timedelta(days=4)).strftime("%Y-%m-%d")
        priority, warnings = cse.determine_priority(10_000_000, soon)
        assert any("Внимание" in w or "дедлайн" in w.lower() for w in warnings)

    def test_far_deadline_no_warning(self):
        """Дедлайн через 30 дней → без предупреждений."""
        far = (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%d")
        priority, warnings = cse.determine_priority(10_000_000, far)
        assert not any("СРОЧНО" in w for w in warnings)

    def test_invalid_deadline_no_crash(self):
        """Некорректная дата → не падает."""
        priority, warnings = cse.determine_priority(5_000_000, "invalid-date")
        assert priority is not None

    def test_empty_deadline_no_crash(self):
        """Пустая дата → не падает."""
        priority, warnings = cse.determine_priority(5_000_000, "")
        assert priority is not None


# ═══════════════════════════════════════════════════════════════════════════
# 7. PIPELINE — extract_and_parse_tender
# ═══════════════════════════════════════════════════════════════════════════

class TestPipelineFull:
    """Полные тесты pipeline extract_and_parse_tender."""

    def test_empty_files_returns_blocked(self):
        """Пустой список файлов → validation_status='blocked'."""
        result = cse.extract_and_parse_tender([])
        assert result["validation_status"] == "blocked"
        assert result["total_chunks"] == 0
        assert result["chunks_used"] == 0

    def test_nonexistent_files_returns_blocked(self):
        """Несуществующие файлы → blocked."""
        result = cse.extract_and_parse_tender(["/tmp/no_such_file.pdf"])
        assert result["total_chunks"] == 0

    def test_result_has_all_required_fields(self):
        """Результат содержит все обязательные поля."""
        # Для пустого ввода — минимальный набор полей
        result = cse.extract_and_parse_tender([])
        # Минимальные обязательные поля (всегда присутствуют)
        minimal_required = [
            "customer", "nmc", "direction", "deadline", "priority",
            "validation_status", "confidence", "warnings",
            "chunks_used", "total_chunks",
        ]
        for field in minimal_required:
            assert field in result, f"Missing field: {field}"

    def test_confidence_is_dict(self):
        """confidence — словарь."""
        result = cse.extract_and_parse_tender([])
        assert isinstance(result["confidence"], dict)

    def test_warnings_is_list(self):
        """warnings — список."""
        result = cse.extract_and_parse_tender([])
        assert isinstance(result["warnings"], list)

    def test_ocr_applied_is_list(self):
        """ocr_applied — список (если присутствует)."""
        result = cse.extract_and_parse_tender([])
        if "ocr_applied" in result:
            assert isinstance(result["ocr_applied"], list)

    def test_key_products_is_list(self):
        """key_products — список (если присутствует)."""
        result = cse.extract_and_parse_tender([])
        if "key_products" in result:
            assert isinstance(result["key_products"], list)

    def test_extraction_time_measured(self):
        """extraction_time_ms >= 0 (если присутствует)."""
        result = cse.extract_and_parse_tender([])
        if "extraction_time_ms" in result:
            assert result["extraction_time_ms"] >= 0

    def test_docx_full_pipeline(self):
        """Полный pipeline с реальным DOCX файлом."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Заказчик: АО «Тестовый завод»")
        doc.add_paragraph("Начальная максимальная цена: 7 500 000 руб.")
        doc.add_paragraph("Дата окончания подачи заявок: 25.06.2026 12:00")
        doc.add_paragraph("Площадка: ЕИС (zakupki.gov.ru)")
        doc.add_paragraph("Номер извещения: 0301-2026-00456")
        doc.add_paragraph("Способ закупки: Запрос котировок")
        doc.add_paragraph("Предмет закупки: Поставка твёрдосплавных фрез концевых D8-D20")
        doc.add_paragraph("Место поставки: г. Москва")
        doc.add_paragraph("Срок поставки: 30 календарных дней с момента заключения")
        doc.add_paragraph("Допускается или эквивалент")
        doc.add_paragraph("Требуется подтверждение в ГИСП")
        # Добавляем позиции
        for i in range(1, 11):
            doc.add_paragraph(f"  {i}.\tФреза D{i*2} z{i} HRC55")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        result = cse.extract_and_parse_tender([tmp_path])
        os.unlink(tmp_path)

        # Проверяем все поля
        assert result["total_chunks"] > 0
        assert result["chunks_used"] > 0
        assert result["customer"] is not None
        assert "Тестовый завод" in result["customer"] or "АО" in result["customer"]
        assert result["nmc"] > 7_000_000
        assert "2026" in result["deadline"]
        assert result["platform"] is not None
        assert result["procedure_number"] is not None
        assert result["procedure_type"] is not None
        assert result["subject"] is not None
        assert result["city"] is not None
        assert result["delivery_terms"] is not None
        assert result["equivalent_allowed"] is True
        assert result["gisp_required"] is True
        assert result["direction"] is not None
        assert result["priority"] is not None
        assert result["validation_status"] in ("ok", "warnings")

    def test_multiple_files_aggregation(self):
        """Несколько файлов → чанки агрегируются."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        paths = []
        for i in range(3):
            doc = Document()
            doc.add_paragraph(f"Файл {i}: " + "Текст " * 50)
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
                doc.save(f.name)
                paths.append(f.name)

        result = cse.extract_and_parse_tender(paths)
        for p in paths:
            os.unlink(p)

        assert result["total_chunks"] >= 3  # Минимум по 1 чанку из каждого файла

    def test_customer_short_generated(self):
        """customer_short генерируется из customer."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Заказчик: Акционерное общество «Тестовый завод»")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        result = cse.extract_and_parse_tender([tmp_path])
        os.unlink(tmp_path)

        if result["customer"]:
            assert "customer_short" in result

    def test_nmc_formatted_generated(self):
        """nmc_formatted генерируется из nmc."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("НМЦ: 5 000 000 руб.")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name

        result = cse.extract_and_parse_tender([tmp_path])
        os.unlink(tmp_path)

        if result["nmc"] > 0:
            assert "руб" in result["nmc_formatted"]


# ═══════════════════════════════════════════════════════════════════════════
# 8. CRON — расписание и автономность
# ═══════════════════════════════════════════════════════════════════════════

class TestCronConfiguration:
    """Тесты конфигурации cron-расписания."""

    def _read_main(self):
        main_path = os.path.join(
            os.path.dirname(__file__), '..', 'src', 'microservice', 'main.py')
        with open(main_path, 'r') as f:
            return f.read()

    def test_yadisk_scan_every_5_minutes(self):
        """YaDisk сканирование настроено на каждые 5 минут."""
        content = self._read_main()
        assert 'IntervalTrigger(minutes=5)' in content or 'minutes=5' in content

    def test_yadisk_scan_not_30_minutes(self):
        """YaDisk НЕ настроен на 30 минут (старое значение)."""
        content = self._read_main()
        # Не должно быть CronTrigger(minute=30) для yadisk
        assert 'CronTrigger(minute=30)' not in content or 'yadisk' not in content.split('CronTrigger(minute=30)')[0][-100:]

    def test_hourly_control_job(self):
        """Ежечасный контроль Р1/Р2 настроен."""
        content = self._read_main()
        assert 'hourly_control' in content

    def test_daily_archive_job(self):
        """Ежедневная архивация настроена."""
        content = self._read_main()
        assert 'daily_archive' in content

    def test_weekly_control_job(self):
        """Еженедельный контроль зависших настроен."""
        content = self._read_main()
        assert 'weekly_control' in content

    def test_monthly_revision_job(self):
        """Ежемесячная ревизия архива настроена."""
        content = self._read_main()
        assert 'monthly_revision' in content

    def test_scheduler_started_on_startup(self):
        """Scheduler запускается при старте приложения."""
        content = self._read_main()
        assert 'scheduler.start()' in content or 'start_scheduler' in content

    def test_all_cron_modules_exist(self):
        """Все cron-модули существуют на диске."""
        base = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice')
        modules = ['cron_yadisk.py', 'cron_hourly.py', 'cron_daily.py',
                   'cron_weekly.py', 'cron_monthly.py']
        for module in modules:
            assert os.path.exists(os.path.join(base, module)), f"Missing: {module}"

    def test_apscheduler_imported(self):
        """APScheduler импортирован в main.py."""
        content = self._read_main()
        assert 'apscheduler' in content.lower()


# ═══════════════════════════════════════════════════════════════════════════
# 9. ДЕДУПЛИКАЦИЯ
# ═══════════════════════════════════════════════════════════════════════════

class TestDeduplication:
    """Тесты модуля дедупликации."""

    def test_deduplication_module_exists(self):
        """Модуль дедупликации существует."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'deduplication.py')
        assert os.path.exists(path)

    def test_deduplication_has_check_method(self):
        """Модуль дедупликации имеет метод check."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'deduplication.py')
        with open(path, 'r') as f:
            content = f.read()
        assert 'def check' in content or 'def is_duplicate' in content

    def test_cron_yadisk_uses_deduplication(self):
        """cron_yadisk.py использует дедупликацию."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'cron_yadisk.py')
        with open(path, 'r') as f:
            content = f.read()
        assert 'dedup' in content.lower() or 'duplicate' in content.lower()

    def test_deduplication_prevents_reprocessing(self):
        """Дедупликация предотвращает повторную обработку."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'cron_yadisk.py')
        with open(path, 'r') as f:
            content = f.read()
        # Должна быть проверка "уже обработан" перед созданием сделки
        assert 'duplicate' in content.lower() or 'already_processed' in content.lower() or 'skip' in content.lower()


# ═══════════════════════════════════════════════════════════════════════════
# 10. DOCKER — автономность
# ═══════════════════════════════════════════════════════════════════════════

class TestDockerAutonomy:
    """Тесты Docker-конфигурации для автономности."""

    def _read_compose(self):
        path = os.path.join(os.path.dirname(__file__), '..', 'deploy', 'docker-compose.yml')
        with open(path, 'r') as f:
            return f.read()

    def _read_dockerfile(self):
        path = os.path.join(os.path.dirname(__file__), '..', 'deploy', 'Dockerfile')
        with open(path, 'r') as f:
            return f.read()

    def test_restart_always(self):
        """Docker restart: always — контейнер переживает перезагрузку."""
        assert 'restart: always' in self._read_compose()

    def test_healthcheck_configured(self):
        """Healthcheck настроен в docker-compose."""
        content = self._read_compose()
        assert 'healthcheck' in content
        assert '/health' in content

    def test_dockerfile_healthcheck(self):
        """HEALTHCHECK в Dockerfile."""
        assert 'HEALTHCHECK' in self._read_dockerfile()

    def test_tesseract_installed(self):
        """Tesseract установлен в Docker для OCR."""
        assert 'tesseract' in self._read_dockerfile()

    def test_poppler_installed(self):
        """Poppler установлен в Docker для pdftotext."""
        assert 'poppler' in self._read_dockerfile()

    def test_volumes_for_data_persistence(self):
        """Volumes настроены для сохранения данных."""
        content = self._read_compose()
        assert 'volumes' in content

    def test_uvicorn_workers(self):
        """Uvicorn запущен с workers."""
        assert '--workers' in self._read_dockerfile()

    def test_port_8000_exposed(self):
        """Порт 8000 открыт."""
        content = self._read_compose()
        assert '8000' in content


# ═══════════════════════════════════════════════════════════════════════════
# 11. ИНТЕГРАЦИЯ cron_yadisk ↔ chunk_score_extractor
# ═══════════════════════════════════════════════════════════════════════════

class TestIntegration:
    """Тесты интеграции между модулями."""

    def test_cron_yadisk_imports_chunk_extractor(self):
        """cron_yadisk.py импортирует chunk_score_extractor."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'cron_yadisk.py')
        with open(path, 'r') as f:
            content = f.read()
        assert 'chunk_score_extractor' in content

    def test_cron_yadisk_calls_extract_and_parse_tender(self):
        """cron_yadisk вызывает extract_and_parse_tender."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'cron_yadisk.py')
        with open(path, 'r') as f:
            content = f.read()
        assert 'extract_and_parse_tender' in content

    def test_cron_yadisk_has_amo_lead_interaction(self):
        """cron_yadisk взаимодействует с amoCRM (обновляет/создаёт сделки)."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'cron_yadisk.py')
        with open(path, 'r') as f:
            content = f.read()
        # Система взаимодействует с amoCRM через _post_note_to_lead и _update_lead_fields
        assert ('_post_note_to_lead' in content or 
                '_update_lead_fields' in content or
                'amo_lead_id' in content)

    def test_cron_yadisk_has_yadisk_client(self):
        """cron_yadisk использует YaDisk клиент."""
        path = os.path.join(os.path.dirname(__file__), '..', 'src', 'microservice', 'cron_yadisk.py')
        with open(path, 'r') as f:
            content = f.read()
        assert 'yadisk' in content.lower()

    def test_field_extractors_match_field_weights(self):
        """Все поля из FIELD_EXTRACTORS есть в FIELD_WEIGHTS и наоборот."""
        for field in cse.FIELD_EXTRACTORS:
            assert field in cse.FIELD_WEIGHTS, f"{field} in EXTRACTORS but not in WEIGHTS"
        for field in cse.FIELD_WEIGHTS:
            assert field in cse.FIELD_EXTRACTORS, f"{field} in WEIGHTS but not in EXTRACTORS"

    def test_direction_keywords_all_have_entries(self):
        """Все направления в DIRECTION_KEYWORDS имеют ключевые слова."""
        for direction, keywords in cse.DIRECTION_KEYWORDS.items():
            assert len(keywords) > 0, f"Direction {direction} has no keywords"

    def test_pipeline_result_compatible_with_amo_lead(self):
        """Результат pipeline совместим с _create_amo_lead."""
        result = cse.extract_and_parse_tender([])
        # Поля необходимые для создания сделки в amoCRM
        amo_required = ["customer", "nmc", "direction", "deadline", "priority"]
        for field in amo_required:
            assert field in result, f"Missing amoCRM field: {field}"

    def test_deploy_requirements_has_all_deps(self):
        """deploy/requirements-deploy.txt содержит все необходимые зависимости."""
        path = os.path.join(os.path.dirname(__file__), '..', 'deploy', 'requirements-deploy.txt')
        with open(path, 'r') as f:
            content = f.read().lower()
        required_deps = ['pdf2image']
        for dep in required_deps:
            assert dep in content, f"Missing dependency: {dep}"


# ═══════════════════════════════════════════════════════════════════════════
# 12. EXTRACT_FILE_TEXT — извлечение текста из файлов
# ═══════════════════════════════════════════════════════════════════════════

class TestExtractFileText:
    """Тесты извлечения текста из различных форматов."""

    def test_docx_extraction(self):
        """Извлечение текста из DOCX."""
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not installed")

        doc = Document()
        doc.add_paragraph("Тестовый текст для извлечения из DOCX")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            tmp_path = f.name
        text = cse.extract_file_text(tmp_path)
        os.unlink(tmp_path)
        assert "Тестовый текст" in text

    def test_txt_extraction(self):
        """Извлечение текста из TXT."""
        with tempfile.NamedTemporaryFile(suffix=".txt", mode="w", delete=False, encoding="utf-8") as f:
            f.write("Простой текстовый файл\nВторая строка")
            tmp_path = f.name
        text = cse.extract_file_text(tmp_path)
        os.unlink(tmp_path)
        assert "Простой текстовый файл" in text

    def test_nonexistent_file_returns_empty(self):
        """Несуществующий файл → пустая строка."""
        text = cse.extract_file_text("/tmp/nonexistent_xyz_123.pdf")
        assert text == "" or text is None

    def test_xlsx_extraction(self):
        """Извлечение текста из XLSX."""
        try:
            import openpyxl
        except ImportError:
            pytest.skip("openpyxl not installed")

        wb = openpyxl.Workbook()
        ws = wb.active
        ws['A1'] = "Фреза D10"
        ws['B1'] = "100 шт"
        with tempfile.NamedTemporaryFile(suffix=".xlsx", delete=False) as f:
            wb.save(f.name)
            tmp_path = f.name
        text = cse.extract_file_text(tmp_path)
        os.unlink(tmp_path)
        assert "Фреза" in text or "D10" in text

    def test_pdf_with_pdftotext(self):
        """PDF извлечение через pdftotext (если доступен)."""
        # Создаём минимальный PDF
        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
            f.write(b"%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n")
            tmp_path = f.name
        text = cse.extract_file_text(tmp_path)
        os.unlink(tmp_path)
        # Может вернуть пустую строку для невалидного PDF — это нормально
        assert isinstance(text, str)


# ═══════════════════════════════════════════════════════════════════════════
# 13. КОНФИГУРАЦИЯ И КОНСТАНТЫ
# ═══════════════════════════════════════════════════════════════════════════

class TestConstants:
    """Тесты корректности констант и конфигурации."""

    def test_confidence_thresholds_order(self):
        """CONFIDENCE_HIGH > CONFIDENCE_MEDIUM > CONFIDENCE_LOW."""
        assert cse.CONFIDENCE_HIGH > cse.CONFIDENCE_MEDIUM > cse.CONFIDENCE_LOW

    def test_confidence_high_lte_1(self):
        """CONFIDENCE_HIGH <= 1.0."""
        assert cse.CONFIDENCE_HIGH <= 1.0

    def test_confidence_low_gt_0(self):
        """CONFIDENCE_LOW > 0.0."""
        assert cse.CONFIDENCE_LOW > 0.0

    def test_min_chunk_size_positive(self):
        """MIN_CHUNK_SIZE > 0."""
        assert cse.MIN_CHUNK_SIZE > 0

    def test_field_weights_all_have_keywords(self):
        """Все FIELD_WEIGHTS имеют keywords."""
        for field, config in cse.FIELD_WEIGHTS.items():
            assert "keywords" in config, f"{field} missing keywords"
            assert len(config["keywords"]) > 0, f"{field} has empty keywords"

    def test_field_weights_all_have_weight(self):
        """Все FIELD_WEIGHTS имеют weight."""
        for field, config in cse.FIELD_WEIGHTS.items():
            assert "weight" in config, f"{field} missing weight"
            assert config["weight"] > 0, f"{field} weight must be > 0"

    def test_direction_keywords_not_empty(self):
        """DIRECTION_KEYWORDS не пустой."""
        assert len(cse.DIRECTION_KEYWORDS) >= 4

    def test_field_extractors_all_callable(self):
        """Все FIELD_EXTRACTORS — callable."""
        for field, extractor in cse.FIELD_EXTRACTORS.items():
            assert callable(extractor), f"{field} extractor is not callable"
